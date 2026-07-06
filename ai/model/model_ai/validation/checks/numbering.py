"""Numbering checks: page number format, location, and section start validation. Keyword: automated document validation"""
from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from model_ai.extractor.models import DocumentMetadata
from model_ai.validation.models import ValidationCheckResult, ValidationIssue
from model_ai.validation.validocx_adapter import _heading_level_from_style

from ._shared import (
    _BAB_RE,
    _FORMAT_ALIAS,
    _NUM_FORMAT_DISPLAY,
    _HEADING_TITLE_MAP_INV,
    _build_occurrences,
)


def _get_page_number_format(sectPr) -> str:
    """Ambil format nomor halaman dari elemen sectPr (w:pgNumType w:fmt).

    Mengembalikan 'decimal' sebagai default apabila elemen w:pgNumType tidak
    ada atau atribut w:fmt tidak di-set — sesuai perilaku default Microsoft
    Word (OOXML spec: nilai default w:fmt adalah 'decimal').
    """
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is not None:
        fmt = pgNumType.get(qn('w:fmt'))
        return fmt if fmt else "decimal"
    return "decimal"


def _hdrftr_has_page_field(hdrftr) -> bool:
    """Cek apakah header/footer mengandung field PAGE."""
    if hdrftr is None:
        return False
    for instrText in hdrftr._element.iter(qn('w:instrText')):
        if 'PAGE' in (instrText.text or '').upper():
            return True
    return False


def _hdrftr_page_alignment(hdrftr) -> str | None:
    """Ambil alignment paragraf yang mengandung field PAGE di header/footer.

    Returns 'LEFT', 'CENTER', 'RIGHT', atau None jika tidak ditemukan.
    """
    if hdrftr is None:
        return None
    _ALIGN_INT_TO_STR = {0: "LEFT", 1: "CENTER", 2: "RIGHT", 3: "JUSTIFY"}
    try:
        for para in hdrftr.paragraphs:
            has_page = any(
                'PAGE' in (el.text or '').upper()
                for el in para._element.iter(qn('w:instrText'))
            )
            if not has_page:
                continue
            align = para.paragraph_format.alignment
            if align is None:
                try:
                    align = para.style.paragraph_format.alignment
                except Exception:
                    pass
            if align is not None:
                return _ALIGN_INT_TO_STR.get(int(align), str(int(align)))
            return "LEFT"
    except Exception:
        pass
    return None


def _read_section_page_numbering(sectPr) -> tuple[str, int]:
    """Baca format dan nomor awal dari elemen w:pgNumType dalam sectPr.

    Returns:
        (fmt, start_num)
        fmt       → "lowerRoman", "upperRoman", "decimal", dsb.
        start_num → nomor halaman awal section (default 1)
    """
    pgNumType = sectPr.find(qn("w:pgNumType"))
    fmt       = "decimal"
    start_num = 1
    if pgNumType is not None:
        fmt = pgNumType.get(qn("w:fmt"), "decimal") or "decimal"
        s   = pgNumType.get(qn("w:start"), "1")
        try:
            start_num = int(s)
        except (TypeError, ValueError):
            start_num = 1
    return fmt, start_num


def _classify_sections_by_metadata(
    doc,
    metadata: "DocumentMetadata",
) -> dict[str, dict | None]:
    """Klasifikasi section dokumen ke zone preliminary/content berdasarkan metadata.

    Menggunakan format penomoran yang diexpect dari metadata (misal lowerRoman vs
    decimal) untuk mencocokkan section DOCX ke zone yang tepat — lebih andal
    daripada mengandalkan posisi heading BAB 1 sebagai satu-satunya penanda.

    Strategi:
    1. Kumpulkan semua sectPr + batas paragraf dari dokumen (inline + body level)
    2. Untuk setiap section, baca pgNumType (fmt + start_num) dan info header/footer
    3. Cocokkan ke zone preliminary/content berdasarkan format dari metadata
    4. Tiebreaker jika ada dua section dengan format sama: section yang mengandung
       heading BAB = content zone

    Fallback jika metadata.numbering kosong:
    - Section dengan fmt romawi (lowerRoman/upperRoman) → preliminary
    - Section dengan fmt decimal → content

    Returns:
        {
            "preliminary": info_dict | None,
            "content":     info_dict | None,
        }
        di mana info_dict memiliki keys:
            fmt, start_num, location, has_header_page, has_footer_page,
            has_any_page, start_para_idx, end_para_idx
    """
    num = metadata.numbering if metadata else None
    prelim_fmt_exp  = (num.preliminary.format if num and num.preliminary else None)
    content_fmt_exp = (num.content.format     if num and num.content     else None)

    body      = doc.element.body
    para_list = list(doc.paragraphs)
    if not para_list:
        return {"preliminary": None, "content": None}

    para_idx_by_id = {id(p._p): i for i, p in enumerate(para_list)}

    raw_sections: list[dict] = []
    prev_end = 0

    for child in body:
        if not (child.tag.endswith('}p') or child.tag == 'p'):
            continue
        pPr = child.find(qn('w:pPr'))
        if pPr is None:
            continue
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is None:
            continue
        para_idx = para_idx_by_id.get(id(child))
        if para_idx is None:
            continue
        raw_sections.append({
            "start_para_idx": prev_end,
            "end_para_idx":   para_idx,
            "sectPr":         sectPr,
        })
        prev_end = para_idx + 1

    body_sectPr = body.find(qn('w:sectPr'))
    if body_sectPr is not None:
        raw_sections.append({
            "start_para_idx": prev_end,
            "end_para_idx":   len(para_list) - 1,
            "sectPr":         body_sectPr,
        })

    if not raw_sections:
        return {"preliminary": None, "content": None}

    def _sec_info(sec: dict) -> dict:
        sp = sec["sectPr"]
        fmt, start_num = _read_section_page_numbering(sp)
        has_own_hdr = bool(sp.findall(qn('w:headerReference')))
        has_own_ftr = bool(sp.findall(qn('w:footerReference')))
        has_header_page = False
        has_footer_page = False
        for s in doc.sections:
            if s._sectPr is sp:
                if has_own_hdr:
                    try:
                        has_header_page = _hdrftr_has_page_field(s.header)
                    except Exception:
                        pass
                if has_own_ftr:
                    try:
                        has_footer_page = _hdrftr_has_page_field(s.footer)
                    except Exception:
                        pass
                break
        location: str | None = (
            "HEADER" if has_header_page else
            "FOOTER" if has_footer_page else None
        )
        return {
            "sectPr":          sp,
            "fmt":             fmt,
            "start_num":       start_num,
            "location":        location,
            "has_header_page": has_header_page,
            "has_footer_page": has_footer_page,
            "has_any_page":    has_header_page or has_footer_page,
            "start_para_idx":  sec["start_para_idx"],
            "end_para_idx":    sec["end_para_idx"],
        }

    section_infos = [_sec_info(s) for s in raw_sections]

    def _norm_fmt(fmt: str | None) -> str | None:
        if not fmt:
            return fmt
        return _FORMAT_ALIAS.get(fmt.lower(), fmt)

    prelim_fmt_exp  = _norm_fmt(prelim_fmt_exp)
    content_fmt_exp = _norm_fmt(content_fmt_exp)

    _HEADING_STYLE_KW = ("heading", "judul", "bab")
    bab1_para_idx: int | None = None
    for i, para in enumerate(para_list):
        style_val  = (para.style.name or "").lower()
        text_upper = (para.text or "").strip().upper()
        if not text_upper.startswith("BAB"):
            continue
        if any(k in style_val for k in _HEADING_STYLE_KW):
            bab1_para_idx = i
            break

    prelim_info:  dict | None = None
    content_info: dict | None = None

    prelim_candidates:  list[dict] = []
    content_candidates: list[dict] = []

    for info in section_infos:
        fmt = _norm_fmt(info["fmt"])
        if prelim_fmt_exp and fmt == prelim_fmt_exp:
            prelim_candidates.append(info)
        elif content_fmt_exp and fmt == content_fmt_exp:
            content_candidates.append(info)
        else:
            if not prelim_fmt_exp and fmt in ("lowerRoman", "upperRoman"):
                prelim_candidates.append(info)
            elif not content_fmt_exp and fmt == "decimal":
                content_candidates.append(info)

    prelim_info = prelim_candidates[0] if prelim_candidates else None

    if content_candidates:
        if bab1_para_idx is not None:
            bab_match = next(
                (i for i in content_candidates
                 if i["start_para_idx"] <= bab1_para_idx <= i["end_para_idx"]),
                None,
            )
            content_info = bab_match or content_candidates[0]
        else:
            content_info = content_candidates[0]

    if content_info is None and bab1_para_idx is not None:
        for info in section_infos:
            if info["start_para_idx"] <= bab1_para_idx <= info["end_para_idx"]:
                content_info = info
                break

    if content_info is None and len(doc.sections) >= 1:
        last_sp = doc.sections[-1]._sectPr
        existing = next(
            (info for info in section_infos if info["sectPr"] is last_sp),
            None,
        )
        if existing is not None:
            content_info = {k: v for k, v in existing.items() if k != "sectPr"}
            content_info["fmt"] = _norm_fmt(existing["fmt"]) or existing["fmt"]
        else:
            fmt_last, start_last = _read_section_page_numbering(last_sp)
            prev_end_idx = (section_infos[-1]["end_para_idx"] + 1
                            if section_infos else 0)
            content_info = {
                "fmt":             _norm_fmt(fmt_last) or fmt_last,
                "start_num":       start_last,
                "location":        None,
                "has_header_page": False,
                "has_footer_page": False,
                "has_any_page":    False,
                "start_para_idx":  prev_end_idx,
                "end_para_idx":    len(para_list) - 1,
            }
            try:
                last_sec = doc.sections[-1]
                has_own_hdr = bool(last_sp.findall(qn('w:headerReference')))
                has_own_ftr = bool(last_sp.findall(qn('w:footerReference')))
                if has_own_hdr:
                    content_info["has_header_page"] = _hdrftr_has_page_field(last_sec.header)
                if has_own_ftr:
                    content_info["has_footer_page"] = _hdrftr_has_page_field(last_sec.footer)
                content_info["has_any_page"] = (
                    content_info["has_header_page"] or content_info["has_footer_page"]
                )
                content_info["location"] = (
                    "HEADER" if content_info["has_header_page"] else
                    "FOOTER" if content_info["has_footer_page"] else None
                )
            except Exception:
                pass

    if prelim_info is None and content_info is not None and len(section_infos) > 1:
        for info in section_infos:
            if info["end_para_idx"] < content_info["start_para_idx"]:
                prelim_info = info

    return {"preliminary": prelim_info, "content": content_info}


def _check_start_section(
    start_at: str,
    doc,
    issues: list[ValidationIssue],
    checks: list[ValidationCheckResult],
    zone: str,
) -> None:
    """Verifikasi bahwa titik mulai penomoran (start_at_section) ada di dokumen."""
    import re
    bab_m = re.match(r'^bab_(\d+)$', start_at, re.IGNORECASE)
    if bab_m:
        target_num = int(bab_m.group(1))
        found = any(
            _heading_level_from_style(para.style) is not None
            and (m := _BAB_RE.match(para.text.strip().upper())) is not None
            and int(m.group(1)) == target_num
            for para in doc.paragraphs
            if para.text.strip()
        )
        label = f"BAB {target_num}"
    else:
        expected_title = _HEADING_TITLE_MAP_INV.get(start_at, start_at.upper().replace("_", " "))
        found = any(
            _heading_level_from_style(para.style) is not None
            and para.text.strip().upper() == expected_title
            for para in doc.paragraphs
        )
        label = expected_title

    field = f"{zone}_start"
    if found:
        checks.append(ValidationCheckResult(
            category="numbering", field=field,
            status="passed",
            message=f"Titik mulai nomor halaman {zone}: '{label}' ditemukan di dokumen",
            expected=start_at,
            occurrences=_build_occurrences([{"text": label, "full_text": label, "style": "Heading 1", "page": None, "bab": None, "para_idx": None}]),
        ))
    else:
        fail_msg = f"Titik mulai nomor halaman {zone} '{label}' tidak ditemukan di dokumen"
        occ_start = _build_occurrences(
            [{"text": label, "full_text": fail_msg, "style": "",
              "page": None, "bab": None, "para_idx": None,
              "actual": "Tidak ditemukan"}],
            actual_str="Tidak ditemukan", expected_str=label,
        ) or None
        issues.append(ValidationIssue(
            category="numbering", field=field,
            severity="error",
            message=fail_msg,
            expected=start_at,
            actual="Tidak ditemukan",
            occurrences=occ_start,
        ))
        checks.append(ValidationCheckResult(
            category="numbering", field=field,
            status="failed",
            message=fail_msg,
            expected=start_at,
            actual="Tidak ditemukan",
            occurrences=occ_start,
        ))


def _check_numbering(
    docx_path: Path,
    metadata: DocumentMetadata,
    doc: DocxDocument | None = None,
) -> tuple[list[ValidationIssue], list[ValidationCheckResult]]:
    """Validasi format dan posisi nomor halaman (preliminary vs content)."""
    issues: list[ValidationIssue] = []
    checks: list[ValidationCheckResult] = []

    num = metadata.numbering
    if num is None:
        checks.append(ValidationCheckResult(
            category="numbering", field="page_number",
            status="skipped",
            message="Tidak ada data numbering di metadata",
            skip_reason="Tidak ada nilai di metadata",
        ))
        return issues, checks

    prelim = num.preliminary
    content = num.content

    if prelim is None and content is None:
        checks.append(ValidationCheckResult(
            category="numbering", field="page_number",
            status="skipped",
            message="Data nomor halaman awal (romawi) dan isi (angka arab) kosong di metadata",
            skip_reason="Tidak ada nilai di metadata",
        ))
        return issues, checks

    try:
        doc = doc or DocxDocument(str(docx_path))

        zones        = _classify_sections_by_metadata(doc, metadata)
        prelim_zone  = zones["preliminary"]
        content_zone = zones["content"]

        all_formats: set[str] = set()
        if prelim_zone:
            all_formats.add(prelim_zone["fmt"])
        if content_zone:
            all_formats.add(content_zone["fmt"])

        if prelim:
            exp_fmt = _FORMAT_ALIAS.get((prelim.format or "").lower(), prelim.format)
            exp_loc = (prelim.location or "").upper()
            exp_start = prelim.start_at_section

            zone = prelim_zone
            fmt_match = zone is not None and zone["fmt"] == exp_fmt

            if fmt_match:
                checks.append(ValidationCheckResult(
                    category="numbering", field="preliminary_format",
                    status="passed",
                    message=(
                        f"Format nomor halaman awal '{exp_fmt}' "
                        f"({_NUM_FORMAT_DISPLAY.get(exp_fmt, exp_fmt)}): sesuai"
                    ),
                    expected=exp_fmt,
                    occurrences=_build_occurrences([{"text": exp_fmt, "full_text": _NUM_FORMAT_DISPLAY.get(exp_fmt, exp_fmt), "style": "", "page": None, "bab": None, "para_idx": None}]),
                ))
                if exp_loc in ("HEADER", "FOOTER") and zone["has_any_page"]:
                    loc_ok = (
                        (exp_loc == "HEADER" and zone["has_header_page"])
                        or (exp_loc == "FOOTER" and zone["has_footer_page"])
                    )
                    if not loc_ok:
                        actual_loc = "HEADER" if zone["has_header_page"] else "FOOTER"
                        msg = (
                            f"Nomor halaman awal seharusnya di {exp_loc}, "
                            f"tetapi ditemukan di {actual_loc}."
                        )
                        occ_prelim_loc = _build_occurrences(
                            [{"text": f"nomor halaman di {actual_loc}",
                              "full_text": msg, "style": "",
                              "page": None, "bab": None, "para_idx": None,
                              "actual": actual_loc}],
                            actual_str=actual_loc, expected_str=exp_loc,
                        ) or None
                        issues.append(ValidationIssue(
                            category="numbering", field="preliminary_location",
                            severity="error", message=msg, expected=exp_loc,
                            actual=actual_loc,
                            occurrences=occ_prelim_loc,
                        ))
                        checks.append(ValidationCheckResult(
                            category="numbering", field="preliminary_location",
                            status="failed", message=msg, expected=exp_loc,
                            actual=actual_loc, occurrences=occ_prelim_loc,
                        ))
                    else:
                        checks.append(ValidationCheckResult(
                            category="numbering", field="preliminary_location",
                            status="passed",
                            message=f"Lokasi nomor halaman awal ({exp_loc}): sesuai",
                            expected=exp_loc,
                            occurrences=_build_occurrences([{"text": exp_loc, "full_text": f"Nomor halaman di {exp_loc}", "style": "", "page": None, "bab": None, "para_idx": None}]),
                        ))
            else:
                actual_fmt = zone["fmt"] if zone else None
                found_fmts = sorted(all_formats)
                msg = (
                    f"Format nomor halaman awal '{exp_fmt}' "
                    f"({_NUM_FORMAT_DISPLAY.get(exp_fmt, exp_fmt)}) tidak ditemukan "
                    f"di bagian sebelum BAB 1. "
                    + (f"Format yang ada: {found_fmts}" if found_fmts else "Tidak ada nomor halaman terdeteksi.")
                )
                occ_prelim_fmt = _build_occurrences(
                    [{"text": actual_fmt or "tidak terdeteksi",
                      "full_text": msg, "style": "",
                      "page": None, "bab": None, "para_idx": None,
                      "actual": actual_fmt or "tidak terdeteksi"}],
                    actual_str=actual_fmt or "tidak terdeteksi", expected_str=exp_fmt,
                ) or None
                issues.append(ValidationIssue(
                    category="numbering", field="preliminary_format",
                    severity="error", message=msg,
                    expected=exp_fmt,
                    actual=actual_fmt,
                    occurrences=occ_prelim_fmt,
                ))
                checks.append(ValidationCheckResult(
                    category="numbering", field="preliminary_format",
                    status="failed", message=msg,
                    expected=exp_fmt,
                    actual=actual_fmt,
                    occurrences=occ_prelim_fmt,
                ))

            if exp_start:
                _check_start_section(exp_start, doc, issues, checks, zone="awal")

            exp_align = (prelim.alignment or "").upper() if prelim else None
            if exp_align:
                actual_align: str | None = None
                for sec in doc.sections:
                    if zone and sec._sectPr is zone["sectPr"]:
                        hdr = sec.header if zone.get("has_header_page") else sec.footer
                        actual_align = _hdrftr_page_alignment(hdr)
                        break
                if actual_align is not None:
                    if actual_align != exp_align:
                        msg_a = (
                            f"Alignment nomor halaman awal seharusnya {exp_align}, "
                            f"ditemukan {actual_align}."
                        )
                        issues.append(ValidationIssue(
                            category="numbering", field="preliminary_alignment",
                            severity="error", message=msg_a,
                            expected=exp_align, actual=actual_align,
                        ))
                        checks.append(ValidationCheckResult(
                            category="numbering", field="preliminary_alignment",
                            status="failed", message=msg_a,
                            expected=exp_align, actual=actual_align,
                        ))
                    else:
                        checks.append(ValidationCheckResult(
                            category="numbering", field="preliminary_alignment",
                            status="passed",
                            message=f"Alignment nomor halaman awal ({exp_align}): sesuai",
                            expected=exp_align, actual=actual_align,
                        ))

        if content:
            exp_fmt = _FORMAT_ALIAS.get((content.format or "").lower(), content.format)
            exp_loc = (content.location or "").upper()
            exp_start = content.start_at_section

            zone = content_zone
            fmt_match = zone is not None and zone["fmt"] == exp_fmt

            if fmt_match:
                checks.append(ValidationCheckResult(
                    category="numbering", field="content_format",
                    status="passed",
                    message=(
                        f"Format nomor halaman isi '{exp_fmt}' "
                        f"({_NUM_FORMAT_DISPLAY.get(exp_fmt, exp_fmt)}): sesuai "
                        f"(ditemukan mulai BAB 1)"
                    ),
                    expected=exp_fmt,
                    occurrences=_build_occurrences([{"text": exp_fmt, "full_text": _NUM_FORMAT_DISPLAY.get(exp_fmt, exp_fmt), "style": "", "page": None, "bab": None, "para_idx": None}]),
                ))
                if exp_loc in ("HEADER", "FOOTER") and zone["has_any_page"]:
                    loc_ok = (
                        (exp_loc == "HEADER" and zone["has_header_page"])
                        or (exp_loc == "FOOTER" and zone["has_footer_page"])
                    )
                    if not loc_ok:
                        actual_loc = "HEADER" if zone["has_header_page"] else "FOOTER"
                        msg = (
                            f"Nomor halaman isi seharusnya di {exp_loc}, "
                            f"tetapi ditemukan di {actual_loc}."
                        )
                        occ_content_loc = _build_occurrences(
                            [{"text": f"nomor halaman di {actual_loc}",
                              "full_text": msg, "style": "",
                              "page": None, "bab": None, "para_idx": None,
                              "actual": actual_loc}],
                            actual_str=actual_loc, expected_str=exp_loc,
                        ) or None
                        issues.append(ValidationIssue(
                            category="numbering", field="content_location",
                            severity="error", message=msg, expected=exp_loc,
                            actual=actual_loc,
                            occurrences=occ_content_loc,
                        ))
                        checks.append(ValidationCheckResult(
                            category="numbering", field="content_location",
                            status="failed", message=msg, expected=exp_loc,
                            actual=actual_loc, occurrences=occ_content_loc,
                        ))
                    else:
                        checks.append(ValidationCheckResult(
                            category="numbering", field="content_location",
                            status="passed",
                            message=f"Lokasi nomor halaman isi ({exp_loc}): sesuai",
                            expected=exp_loc,
                            occurrences=_build_occurrences([{"text": exp_loc, "full_text": f"Nomor halaman di {exp_loc}", "style": "", "page": None, "bab": None, "para_idx": None}]),
                        ))
            else:
                actual_fmt = zone["fmt"] if zone else None
                found_fmts = sorted(all_formats)
                msg = (
                    f"Format nomor halaman isi '{exp_fmt}' "
                    f"({_NUM_FORMAT_DISPLAY.get(exp_fmt, exp_fmt)}) tidak ditemukan "
                    f"di section yang mengandung BAB 1. "
                    + (f"Format yang ada: {found_fmts}" if found_fmts else "Tidak ada nomor halaman terdeteksi.")
                )
                occ_content_fmt = _build_occurrences(
                    [{"text": actual_fmt or "tidak terdeteksi",
                      "full_text": msg, "style": "",
                      "page": None, "bab": None, "para_idx": None,
                      "actual": actual_fmt or "tidak terdeteksi"}],
                    actual_str=actual_fmt or "tidak terdeteksi", expected_str=exp_fmt,
                ) or None
                issues.append(ValidationIssue(
                    category="numbering", field="content_format",
                    severity="error", message=msg,
                    expected=exp_fmt,
                    actual=actual_fmt,
                    occurrences=occ_content_fmt,
                ))
                checks.append(ValidationCheckResult(
                    category="numbering", field="content_format",
                    status="failed", message=msg,
                    expected=exp_fmt,
                    actual=actual_fmt,
                    occurrences=occ_content_fmt,
                ))

            if exp_start:
                _check_start_section(exp_start, doc, issues, checks, zone="isi")

            exp_align_c = (content.alignment or "").upper() if content else None
            if exp_align_c:
                actual_align_c: str | None = None
                for sec in doc.sections:
                    if content_zone and sec._sectPr is content_zone["sectPr"]:
                        hdr_c = sec.header if content_zone.get("has_header_page") else sec.footer
                        actual_align_c = _hdrftr_page_alignment(hdr_c)
                        break
                if actual_align_c is not None:
                    if actual_align_c != exp_align_c:
                        msg_ca = (
                            f"Alignment nomor halaman isi seharusnya {exp_align_c}, "
                            f"ditemukan {actual_align_c}."
                        )
                        issues.append(ValidationIssue(
                            category="numbering", field="content_alignment",
                            severity="error", message=msg_ca,
                            expected=exp_align_c, actual=actual_align_c,
                        ))
                        checks.append(ValidationCheckResult(
                            category="numbering", field="content_alignment",
                            status="failed", message=msg_ca,
                            expected=exp_align_c, actual=actual_align_c,
                        ))
                    else:
                        checks.append(ValidationCheckResult(
                            category="numbering", field="content_alignment",
                            status="passed",
                            message=f"Alignment nomor halaman isi ({exp_align_c}): sesuai",
                            expected=exp_align_c, actual=actual_align_c,
                        ))

    except Exception as exc:
        checks.append(ValidationCheckResult(
            category="numbering", field="page_number",
            status="skipped",
            message=f"Pengecekan nomor halaman dilewati: {exc}",
            skip_reason=str(exc),
        ))

    return issues, checks
