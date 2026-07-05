"""Figures & tables checks: caption format and positioning. Keyword: automated document validation"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from model_ai.extractor.models import DocumentMetadata
from model_ai.validation.models import ValidationCheckResult, ValidationIssue

from ._shared import (
    _BAB_RE,
    _FIG_DETECT_RE,
    _TBL_DETECT_RE,
    _LAMPIRAN_BROAD_RE,
    _CAPTION_ALIGN_MAP,
    _ALIGN_LABEL,
    _build_occurrences,
)


def _template_to_regex(template: str) -> re.Pattern:
    """Konversi template caption seperti 'Gambar {n}. {title}' ke regex.

    Titik (.) yang memisahkan nomor bab/urutan diizinkan diikuti spasi opsional,
    sehingga '4.1.' dan '4. 1.' sama-sama diterima.
    """
    escaped = re.escape(template)
    escaped = escaped.replace(r'\{n\}', r'\d+')
    escaped = escaped.replace(r'\{bab\}', r'\d+')
    escaped = escaped.replace(r'\{title\}', r'.+')
    escaped = escaped.replace(r'\.', r'\.\s*')
    return re.compile(r'^' + escaped, re.IGNORECASE)


def _para_contains_image(para) -> bool:
    """Cek apakah paragraf mengandung gambar inline."""
    el = para._element
    return (
        el.find('.//' + qn('w:drawing')) is not None
        or el.find('.//' + qn('w:pict')) is not None
    )


def _get_page_number_format_for_content(sectPr) -> str:
    """Ambil format nomor halaman dari elemen sectPr (w:pgNumType w:fmt).

    Dipakai oleh _build_content_elements. Mengembalikan 'decimal' sebagai
    default apabila elemen w:pgNumType tidak ada atau atribut w:fmt tidak di-set.
    """
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is not None:
        fmt = pgNumType.get(qn('w:fmt'))
        return fmt if fmt else "decimal"
    return "decimal"


def _build_content_elements(doc) -> tuple[list[tuple[str, object]], str]:
    """Bangun daftar elemen body yang dibatasi pada section dengan penomoran decimal.

    Urutan prioritas batas scan:
      1. Mulai dari awal section yang punya pgNumType decimal
      2. Fallback: mulai dari Heading 1 BAB pertama jika tidak ada section decimal
      3. Berhenti tepat sebelum heading DAFTAR PUSTAKA atau LAMPIRAN

    Returns:
        (elements, source) di mana source menjelaskan metode yang dipakai.
    """
    body = doc.element.body
    para_by_el = {id(p._element): p for p in doc.paragraphs}
    tbl_by_el  = {id(t._element): t for t in doc.tables}

    all_elements: list[tuple[str, object]] = []
    section_ends: list[tuple[int, str | None]] = []

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p' and id(child) in para_by_el:
            para = para_by_el[id(child)]
            idx = len(all_elements)
            all_elements.append(("para", para))
            pPr = child.find(qn('w:pPr'))
            if pPr is not None:
                sectPr = pPr.find(qn('w:sectPr'))
                if sectPr is not None:
                    section_ends.append((idx, _get_page_number_format_for_content(sectPr)))
        elif tag == 'tbl' and id(child) in tbl_by_el:
            all_elements.append(("table", tbl_by_el[id(child)]))

    body_sectPr = body.find(qn('w:sectPr'))
    if body_sectPr is not None:
        section_ends.append((len(all_elements) - 1, _get_page_number_format_for_content(body_sectPr)))

    decimal_start: int | None = None
    decimal_end:   int | None = None
    prev_end = -1
    for end_idx, fmt in section_ends:
        if fmt == "decimal":
            if decimal_start is None:
                decimal_start = prev_end + 1
            decimal_end = end_idx
        prev_end = end_idx

    if decimal_start is not None and decimal_end is not None:
        candidate = all_elements[decimal_start : decimal_end + 1]
        source = "decimal_section"
    else:
        bab1_idx = next(
            (i for i, (etype, elem) in enumerate(all_elements)
             if etype == "para"
             and elem.style.name == "Heading 1"
             and _BAB_RE.match((elem.text or "").strip().upper())),
            0,
        )
        candidate = all_elements[bab1_idx:]
        source = "bab1_fallback"

    _EXCLUDED_HEADINGS = frozenset({"DAFTAR PUSTAKA", "LAMPIRAN"})
    cutoff = len(candidate)
    for i, (etype, elem) in enumerate(candidate):
        if etype == "para" and elem.style.name == "Heading 1":
            if (elem.text or "").strip().upper() in _EXCLUDED_HEADINGS:
                cutoff = i
                break

    return candidate[:cutoff], source


def _check_caption_format(
    docx_path: Path,
    metadata: DocumentMetadata,
    doc: DocxDocument | None = None,
) -> tuple[list[ValidationIssue], list[ValidationCheckResult]]:
    """Validasi atribut caption gambar/tabel via text-pattern, bukan style name.

    Caption dideteksi dari teks yang diawali 'Gambar <angka>' atau 'Tabel <angka>'.
    Alignment dibaca dari metadata.figures_and_tables per tipe caption (CENTER fallback).
    Font family dan font size harus sama dengan body — dicek per tipe caption.
    Style name diabaikan agar tidak false positive pada nama dinamis.
    """
    issues: list[ValidationIssue] = []
    checks: list[ValidationCheckResult] = []

    t  = metadata.typography
    ft = metadata.figures_and_tables

    expected_font = t.font_family if t else None
    expected_size = int(t.font_size_body_pt) if t and t.font_size_body_pt else None

    fig_align_str = ((ft.caption_alignment_figure or "CENTER").upper() if ft else "CENTER")
    tbl_align_str = ((ft.caption_alignment_table  or "CENTER").upper() if ft else "CENTER")
    fig_align_val = _CAPTION_ALIGN_MAP.get(fig_align_str)
    tbl_align_val = _CAPTION_ALIGN_MAP.get(tbl_align_str)

    try:
        doc = doc or DocxDocument(str(docx_path))

        wrong_fig_alignment: list[dict] = []
        wrong_tbl_alignment: list[dict] = []
        wrong_font_items:    list[dict] = []
        wrong_size_items:    list[dict] = []
        fig_total = 0
        tbl_total = 0
        fig_pass_align_items: list[dict] = []
        tbl_pass_align_items: list[dict] = []
        font_pass_items:      list[dict] = []
        size_pass_items:      list[dict] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            is_fig = bool(_FIG_DETECT_RE.match(text))
            is_tbl = bool(_TBL_DETECT_RE.match(text))
            if not is_fig and not is_tbl:
                continue

            if is_fig:
                fig_total += 1
            else:
                tbl_total += 1

            para_info = {"text": text[:100], "full_text": text, "style": para.style.name, "page": None, "bab": None, "para_idx": None}

            align = para.paragraph_format.alignment
            if align is None:
                try:
                    align = para.style.paragraph_format.alignment
                except Exception:
                    align = None

            if align is not None:
                align_label = _ALIGN_LABEL.get(align.value if hasattr(align, "value") else align, str(align))
                if is_fig:
                    if align != fig_align_val:
                        wrong_fig_alignment.append({**para_info, "actual": align_label})
                    else:
                        fig_pass_align_items.append(para_info)
                elif is_tbl:
                    if align != tbl_align_val:
                        wrong_tbl_alignment.append({**para_info, "actual": align_label})
                    else:
                        tbl_pass_align_items.append(para_info)

            for run in para.runs:
                if not run.text.strip():
                    continue

                actual_font: str | None = run.font.name
                if actual_font is None:
                    try:
                        actual_font = para.style.font.name
                    except Exception:
                        pass

                actual_size_pt: int | None = None
                if run.font.size is not None:
                    actual_size_pt = round(run.font.size.pt)
                else:
                    try:
                        if para.style.font.size is not None:
                            actual_size_pt = round(para.style.font.size.pt)
                    except Exception:
                        pass

                if expected_font and actual_font:
                    item = {**para_info, "actual": actual_font}
                    if actual_font != expected_font:
                        wrong_font_items.append(item)
                    else:
                        font_pass_items.append(item)

                if expected_size and actual_size_pt is not None:
                    actual_size_str = f"{actual_size_pt}pt"
                    item = {**para_info, "actual": actual_size_str}
                    if actual_size_pt != expected_size:
                        wrong_size_items.append(item)
                    else:
                        size_pass_items.append(item)

                break

        if fig_total > 0:
            if wrong_fig_alignment:
                first_act = wrong_fig_alignment[0].get("actual", f"bukan {fig_align_str}")
                msg = (
                    f"{len(wrong_fig_alignment)} caption gambar tidak {fig_align_str}. "
                    f'Contoh: "{wrong_fig_alignment[0]["text"]}"'
                )
                _occ_fig_align = _build_occurrences(
                    wrong_fig_alignment, actual_str=None, expected_str=fig_align_str
                ) or None
                issues.append(ValidationIssue(
                    category="figures_tables", field="caption_alignment_figure",
                    severity="error", message=msg,
                    expected=fig_align_str, actual=first_act,
                    occurrences=_occ_fig_align,
                ))
                checks.append(ValidationCheckResult(
                    category="figures_tables", field="caption_alignment_figure",
                    status="failed", message=msg,
                    expected=fig_align_str, actual=first_act,
                    occurrences=_occ_fig_align,
                ))
            else:
                checks.append(ValidationCheckResult(
                    category="figures_tables", field="caption_alignment_figure",
                    status="passed",
                    message=f"Semua {fig_total} caption gambar alignment {fig_align_str}",
                    expected=fig_align_str,
                    occurrences=_build_occurrences(fig_pass_align_items),
                ))

        if tbl_total > 0:
            if wrong_tbl_alignment:
                first_act = wrong_tbl_alignment[0].get("actual", f"bukan {tbl_align_str}")
                msg = (
                    f"{len(wrong_tbl_alignment)} caption tabel tidak {tbl_align_str}. "
                    f'Contoh: "{wrong_tbl_alignment[0]["text"]}"'
                )
                _occ_tbl_align = _build_occurrences(
                    wrong_tbl_alignment, actual_str=None, expected_str=tbl_align_str
                ) or None
                issues.append(ValidationIssue(
                    category="figures_tables", field="caption_alignment_table",
                    severity="error", message=msg,
                    expected=tbl_align_str, actual=first_act,
                    occurrences=_occ_tbl_align,
                ))
                checks.append(ValidationCheckResult(
                    category="figures_tables", field="caption_alignment_table",
                    status="failed", message=msg,
                    expected=tbl_align_str, actual=first_act,
                    occurrences=_occ_tbl_align,
                ))
            else:
                checks.append(ValidationCheckResult(
                    category="figures_tables", field="caption_alignment_table",
                    status="passed",
                    message=f"Semua {tbl_total} caption tabel alignment {tbl_align_str}",
                    expected=tbl_align_str,
                    occurrences=_build_occurrences(tbl_pass_align_items),
                ))

        total_captions = fig_total + tbl_total
        if wrong_font_items:
            first_actual = wrong_font_items[0].get("actual", "")
            msg = (
                f"{len(wrong_font_items)} caption font tidak sesuai "
                f"(seharusnya: {expected_font}). "
                f'Contoh: "{wrong_font_items[0]["text"]}"'
            )
            _occ_font = _build_occurrences(
                wrong_font_items, actual_str=None, expected_str=expected_font
            ) or None
            issues.append(ValidationIssue(
                category="figures_tables", field="caption_font",
                severity="error", message=msg,
                expected=expected_font, actual=first_actual,
                occurrences=_occ_font,
            ))
            checks.append(ValidationCheckResult(
                category="figures_tables", field="caption_font",
                status="failed", message=msg,
                expected=expected_font, actual=first_actual,
                occurrences=_occ_font,
            ))
        elif expected_font and total_captions > 0:
            checks.append(ValidationCheckResult(
                category="figures_tables", field="caption_font",
                status="passed",
                message=f"Font caption sesuai: {expected_font}",
                expected=expected_font,
                occurrences=_build_occurrences(
                    font_pass_items, actual_str=None, expected_str=None
                ),
            ))

        if wrong_size_items:
            first_actual_size = wrong_size_items[0].get("actual", "")
            msg = (
                f"{len(wrong_size_items)} caption ukuran font tidak sesuai "
                f"(seharusnya: {expected_size}pt). "
                f'Contoh: "{wrong_size_items[0]["text"]}"'
            )
            _occ_size = _build_occurrences(
                wrong_size_items, actual_str=None, expected_str=f"{expected_size}pt"
            ) or None
            issues.append(ValidationIssue(
                category="figures_tables", field="caption_font_size",
                severity="error", message=msg,
                expected=f"{expected_size}pt", actual=first_actual_size,
                occurrences=_occ_size,
            ))
            checks.append(ValidationCheckResult(
                category="figures_tables", field="caption_font_size",
                status="failed", message=msg,
                expected=f"{expected_size}pt", actual=first_actual_size,
                occurrences=_occ_size,
            ))
        elif expected_size and total_captions > 0:
            checks.append(ValidationCheckResult(
                category="figures_tables", field="caption_font_size",
                status="passed",
                message=f"Ukuran font caption sesuai: {expected_size}pt",
                expected=f"{expected_size}pt",
                occurrences=_build_occurrences(
                    size_pass_items, actual_str=None, expected_str=None
                ),
            ))

        if total_captions == 0:
            for _fld in ("caption_alignment_figure", "caption_alignment_table"):
                checks.append(ValidationCheckResult(
                    category="figures_tables", field=_fld,
                    status="skipped",
                    message="Tidak ada caption gambar/tabel ditemukan",
                    skip_reason="Tidak ada caption",
                ))

    except Exception as exc:
        checks.append(ValidationCheckResult(
            category="figures_tables", field="caption_alignment_figure",
            status="skipped",
            message=f"Pengecekan atribut caption dilewati: {exc}",
            skip_reason=str(exc),
        ))

    return issues, checks


def _check_figures_tables(
    docx_path: Path,
    metadata: DocumentMetadata,
    doc: DocxDocument | None = None,
) -> tuple[list[ValidationIssue], list[ValidationCheckResult]]:
    """Validasi posisi caption dan format penomoran gambar/tabel + lampiran.

    Kepemilikan field (tidak overlap dengan _check_lampiran_format):
      - figure_caption_position / table_caption_position — posisi relatif gambar/tabel
      - figure_caption_format / table_caption_format — template penomoran (via _build_content_elements)
      - lampiran_caption_format — template penomoran header lampiran (via doc.paragraphs scan terpisah)
      - lampiran_caption_alignment — alignment header lampiran

    _check_lampiran_format() memiliki: lampiran_separator, lampiran_font, lampiran_spacing.
    Keduanya scan _LAMPIRAN_BROAD_RE tetapi mengecek field yang berbeda.
    """
    issues: list[ValidationIssue] = []
    checks: list[ValidationCheckResult] = []

    ft = metadata.figures_and_tables
    if ft is None:
        checks.append(ValidationCheckResult(
            category="figures_tables", field="caption",
            status="skipped",
            message="Tidak ada data figures_and_tables di metadata",
            skip_reason="Tidak ada nilai di metadata",
        ))
        return issues, checks

    tbl_pos_exp  = (ft.table_caption_position or "").upper()
    fig_pos_exp  = (ft.figure_caption_position or "").upper()
    fig_fmt_tpl  = ft.caption_format_figure
    tbl_fmt_tpl  = ft.caption_format_table
    lamp_fmt_tpl = ft.caption_format_lampiran
    lamp_align_str = (ft.caption_alignment_lampiran or "").upper() or None

    if not tbl_pos_exp and not fig_pos_exp and not fig_fmt_tpl and not tbl_fmt_tpl \
            and not lamp_fmt_tpl and not lamp_align_str:
        checks.append(ValidationCheckResult(
            category="figures_tables", field="caption",
            status="skipped",
            message="Tidak ada aturan caption di metadata",
            skip_reason="Tidak ada nilai di metadata",
        ))
        return issues, checks

    try:
        doc = doc or DocxDocument(str(docx_path))

        fig_fmt_re  = _template_to_regex(fig_fmt_tpl)  if fig_fmt_tpl  else None
        tbl_fmt_re  = _template_to_regex(tbl_fmt_tpl)  if tbl_fmt_tpl  else None
        lamp_fmt_re = _template_to_regex(lamp_fmt_tpl) if lamp_fmt_tpl else None
        lamp_align_val = (
            _CAPTION_ALIGN_MAP.get(lamp_align_str)
            if lamp_align_str else None
        )

        elements, scan_source = _build_content_elements(doc)

        fig_pos_errors: list[dict] = []
        fig_fmt_errors: list[dict] = []
        tbl_pos_errors: list[dict] = []
        tbl_fmt_errors: list[dict] = []
        fig_count = 0
        tbl_count = 0
        fig_pos_pass_items: list[dict] = []
        fig_fmt_pass_items: list[dict] = []
        tbl_pos_pass_items: list[dict] = []
        tbl_fmt_pass_items: list[dict] = []

        current_bab: str | None = None
        for i, (etype, elem) in enumerate(elements):
            if etype != "para":
                continue
            text = elem.text.strip() if hasattr(elem, 'text') and elem.text else ""
            if not text:
                continue

            # Lacak BAB aktif untuk konteks lokasi occurrence
            bab_m = _BAB_RE.match(text.upper())
            if bab_m and any(k in (elem.style.name or "").lower() for k in ("heading", "judul")):
                current_bab = f"BAB {bab_m.group(1)}"

            if _FIG_DETECT_RE.match(text):
                fig_count += 1
                fig_para_info = {"text": text[:100], "full_text": text, "style": elem.style.name if elem.style else "", "page": None, "bab": current_bab, "para_idx": None}
                if fig_fmt_re:
                    if not fig_fmt_re.match(text):
                        fig_fmt_errors.append(fig_para_info)
                    else:
                        fig_fmt_pass_items.append(fig_para_info)
                if fig_pos_exp == "BELOW":
                    found_img = any(
                        elements[j][0] == "para" and _para_contains_image(elements[j][1])
                        for j in range(max(0, i - 3), i)
                    )
                    if not found_img:
                        fig_pos_errors.append(fig_para_info)
                    else:
                        fig_pos_pass_items.append(fig_para_info)
                elif fig_pos_exp == "ABOVE":
                    found_img = any(
                        elements[j][0] == "para" and _para_contains_image(elements[j][1])
                        for j in range(i + 1, min(len(elements), i + 4))
                    )
                    if not found_img:
                        fig_pos_errors.append(fig_para_info)
                    else:
                        fig_pos_pass_items.append(fig_para_info)

            elif _TBL_DETECT_RE.match(text):
                tbl_count += 1
                tbl_para_info = {"text": text[:100], "full_text": text, "style": elem.style.name if elem.style else "", "page": None, "bab": current_bab, "para_idx": None}
                if tbl_fmt_re:
                    if not tbl_fmt_re.match(text):
                        tbl_fmt_errors.append(tbl_para_info)
                    else:
                        tbl_fmt_pass_items.append(tbl_para_info)
                if tbl_pos_exp == "ABOVE":
                    next_is_tbl = i + 1 < len(elements) and elements[i + 1][0] == "table"
                    if not next_is_tbl:
                        tbl_pos_errors.append(tbl_para_info)
                    else:
                        tbl_pos_pass_items.append(tbl_para_info)
                elif tbl_pos_exp == "BELOW":
                    prev_is_tbl = i > 0 and elements[i - 1][0] == "table"
                    if not prev_is_tbl:
                        tbl_pos_errors.append(tbl_para_info)
                    else:
                        tbl_pos_pass_items.append(tbl_para_info)

        if fig_count == 0 and tbl_count == 0:
            scan_label = (
                "section dengan nomor halaman angka arab"
                if scan_source == "decimal_section"
                else "mulai BAB 1 (fallback — section decimal tidak ditemukan)"
            )
            checks.append(ValidationCheckResult(
                category="figures_tables", field="caption",
                status="skipped",
                message=f"Tidak ditemukan caption gambar atau tabel di area scan: {scan_label}",
                skip_reason="Tidak ada caption terdeteksi",
            ))
            return issues, checks

        if fig_count > 0:
            if fig_pos_errors:
                msg = (
                    f"Caption gambar seharusnya {fig_pos_exp} gambar. "
                    f"{len(fig_pos_errors)}x salah posisi. "
                    f"Contoh: \"{fig_pos_errors[0]['text']}\""
                )
                occ_fig_pos = _build_occurrences(
                    fig_pos_errors,
                    actual_str=f"bukan {fig_pos_exp}", expected_str=fig_pos_exp,
                ) or None
                issues.append(ValidationIssue(
                    category="figures_tables", field="figure_caption_position",
                    severity="error", message=msg, expected=fig_pos_exp,
                    actual=f"bukan {fig_pos_exp}",
                    occurrences=occ_fig_pos,
                ))
                checks.append(ValidationCheckResult(
                    category="figures_tables", field="figure_caption_position",
                    status="failed", message=msg, expected=fig_pos_exp,
                    actual=f"bukan {fig_pos_exp}", occurrences=occ_fig_pos,
                ))
            else:
                checks.append(ValidationCheckResult(
                    category="figures_tables", field="figure_caption_position",
                    status="passed",
                    message=f"Posisi caption gambar ({fig_pos_exp}): {fig_count} caption sesuai",
                    expected=fig_pos_exp,
                    occurrences=_build_occurrences(fig_pos_pass_items),
                ))

            if fig_fmt_re:
                if fig_fmt_errors:
                    msg = (
                        f"Format caption gambar tidak sesuai pola '{fig_fmt_tpl}'. "
                        f"{len(fig_fmt_errors)}x salah format. "
                        f"Contoh: \"{fig_fmt_errors[0]['text']}\""
                    )
                    occ_fig_fmt = _build_occurrences(
                        fig_fmt_errors,
                        actual_str=None, expected_str=fig_fmt_tpl,
                    ) or None
                    issues.append(ValidationIssue(
                        category="figures_tables", field="figure_caption_format",
                        severity="error", message=msg,
                        expected=fig_fmt_tpl, actual=fig_fmt_errors[0]["text"],
                        occurrences=occ_fig_fmt,
                    ))
                    checks.append(ValidationCheckResult(
                        category="figures_tables", field="figure_caption_format",
                        status="failed", message=msg,
                        expected=fig_fmt_tpl, actual=fig_fmt_errors[0]["text"],
                        occurrences=occ_fig_fmt,
                    ))
                else:
                    checks.append(ValidationCheckResult(
                        category="figures_tables", field="figure_caption_format",
                        status="passed",
                        message=f"Format caption gambar '{fig_fmt_tpl}': {fig_count} caption sesuai",
                        expected=fig_fmt_tpl,
                        occurrences=_build_occurrences(fig_fmt_pass_items),
                    ))

        if tbl_count > 0:
            if tbl_pos_errors:
                msg = (
                    f"Caption tabel seharusnya {tbl_pos_exp} tabel. "
                    f"{len(tbl_pos_errors)}x salah posisi. "
                    f"Contoh: \"{tbl_pos_errors[0]['text']}\""
                )
                occ_tbl_pos = _build_occurrences(
                    tbl_pos_errors,
                    actual_str=f"bukan {tbl_pos_exp}", expected_str=tbl_pos_exp,
                ) or None
                issues.append(ValidationIssue(
                    category="figures_tables", field="table_caption_position",
                    severity="error", message=msg, expected=tbl_pos_exp,
                    actual=f"bukan {tbl_pos_exp}",
                    occurrences=occ_tbl_pos,
                ))
                checks.append(ValidationCheckResult(
                    category="figures_tables", field="table_caption_position",
                    status="failed", message=msg, expected=tbl_pos_exp,
                    actual=f"bukan {tbl_pos_exp}", occurrences=occ_tbl_pos,
                ))
            else:
                checks.append(ValidationCheckResult(
                    category="figures_tables", field="table_caption_position",
                    status="passed",
                    message=f"Posisi caption tabel ({tbl_pos_exp}): {tbl_count} caption sesuai",
                    expected=tbl_pos_exp,
                    occurrences=_build_occurrences(tbl_pos_pass_items),
                ))

            if tbl_fmt_re:
                if tbl_fmt_errors:
                    msg = (
                        f"Format caption tabel tidak sesuai pola '{tbl_fmt_tpl}'. "
                        f"{len(tbl_fmt_errors)}x salah format. "
                        f"Contoh: \"{tbl_fmt_errors[0]['text']}\""
                    )
                    occ_tbl_fmt = _build_occurrences(
                        tbl_fmt_errors,
                        actual_str=None, expected_str=tbl_fmt_tpl,
                    ) or None
                    issues.append(ValidationIssue(
                        category="figures_tables", field="table_caption_format",
                        severity="error", message=msg,
                        expected=tbl_fmt_tpl, actual=tbl_fmt_errors[0]["text"],
                        occurrences=occ_tbl_fmt,
                    ))
                    checks.append(ValidationCheckResult(
                        category="figures_tables", field="table_caption_format",
                        status="failed", message=msg,
                        expected=tbl_fmt_tpl, actual=tbl_fmt_errors[0]["text"],
                        occurrences=occ_tbl_fmt,
                    ))
                else:
                    checks.append(ValidationCheckResult(
                        category="figures_tables", field="table_caption_format",
                        status="passed",
                        message=f"Format caption tabel '{tbl_fmt_tpl}': {tbl_count} caption sesuai",
                        expected=tbl_fmt_tpl,
                        occurrences=_build_occurrences(tbl_fmt_pass_items),
                    ))

        if lamp_fmt_re or lamp_align_val is not None:
            lamp_count             = 0
            lamp_fmt_errors:    list[dict] = []
            lamp_align_errors:  list[dict] = []
            lamp_fmt_pass_items:   list[dict] = []
            lamp_align_pass_items: list[dict] = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text or not _LAMPIRAN_BROAD_RE.match(text):
                    continue
                lamp_count += 1
                lamp_para_info = {"text": text[:100], "full_text": text, "style": para.style.name, "page": None, "bab": None, "para_idx": None}

                if lamp_fmt_re:
                    if not lamp_fmt_re.match(text):
                        lamp_fmt_errors.append(lamp_para_info)
                    else:
                        lamp_fmt_pass_items.append(lamp_para_info)

                if lamp_align_val is not None:
                    align = para.paragraph_format.alignment
                    if align is None:
                        try:
                            align = para.style.paragraph_format.alignment
                        except Exception:
                            align = None
                    if align is not None and align != lamp_align_val:
                        lamp_align_errors.append({**lamp_para_info, "actual": _ALIGN_LABEL.get(int(align), str(align))})
                    else:
                        lamp_align_pass_items.append(lamp_para_info)

            if lamp_fmt_re:
                if lamp_count == 0:
                    checks.append(ValidationCheckResult(
                        category="figures_tables", field="lampiran_caption_format",
                        status="skipped",
                        message="Tidak ditemukan caption lampiran di dokumen",
                        skip_reason="Tidak ada paragraf diawali 'Lampiran '",
                    ))
                elif lamp_fmt_errors:
                    msg = (
                        f"Format caption lampiran tidak sesuai pola '{lamp_fmt_tpl}'. "
                        f"{len(lamp_fmt_errors)}x salah. "
                        f"Contoh: \"{lamp_fmt_errors[0]['text']}\""
                    )
                    occ_lamp_fmt = _build_occurrences(
                        lamp_fmt_errors,
                        actual_str=None, expected_str=lamp_fmt_tpl,
                    ) or None
                    issues.append(ValidationIssue(
                        category="figures_tables", field="lampiran_caption_format",
                        severity="error", message=msg,
                        expected=lamp_fmt_tpl, actual=lamp_fmt_errors[0]["text"],
                        occurrences=occ_lamp_fmt,
                    ))
                    checks.append(ValidationCheckResult(
                        category="figures_tables", field="lampiran_caption_format",
                        status="failed", message=msg,
                        expected=lamp_fmt_tpl, actual=lamp_fmt_errors[0]["text"],
                        occurrences=occ_lamp_fmt,
                    ))
                else:
                    checks.append(ValidationCheckResult(
                        category="figures_tables", field="lampiran_caption_format",
                        status="passed",
                        message=f"Format caption lampiran '{lamp_fmt_tpl}': {lamp_count} caption sesuai",
                        expected=lamp_fmt_tpl,
                        occurrences=_build_occurrences(lamp_fmt_pass_items),
                    ))

            if lamp_align_val is not None:
                if lamp_count == 0:
                    checks.append(ValidationCheckResult(
                        category="figures_tables", field="lampiran_caption_alignment",
                        status="skipped",
                        message="Tidak ditemukan caption lampiran di dokumen",
                        skip_reason="Tidak ada paragraf diawali 'Lampiran '",
                    ))
                elif lamp_align_errors:
                    _first_actual = lamp_align_errors[0].get("actual", f"bukan {lamp_align_str}")
                    msg = (
                        f"{len(lamp_align_errors)} caption lampiran tidak {lamp_align_str}. "
                        f"Contoh: \"{lamp_align_errors[0]['text']}\""
                    )
                    occ_lamp_align = _build_occurrences(
                        lamp_align_errors,
                        actual_str=_first_actual, expected_str=lamp_align_str,
                    ) or None
                    issues.append(ValidationIssue(
                        category="figures_tables", field="lampiran_caption_alignment",
                        severity="error", message=msg,
                        expected=lamp_align_str, actual=_first_actual,
                        occurrences=occ_lamp_align,
                    ))
                    checks.append(ValidationCheckResult(
                        category="figures_tables", field="lampiran_caption_alignment",
                        status="failed", message=msg,
                        expected=lamp_align_str, actual=_first_actual,
                        occurrences=occ_lamp_align,
                    ))
                else:
                    checks.append(ValidationCheckResult(
                        category="figures_tables", field="lampiran_caption_alignment",
                        status="passed",
                        message=f"Semua {lamp_count} caption lampiran alignment {lamp_align_str}",
                        expected=lamp_align_str,
                        occurrences=_build_occurrences(lamp_align_pass_items),
                    ))

    except Exception as exc:
        checks.append(ValidationCheckResult(
            category="figures_tables", field="caption",
            status="skipped",
            message=f"Pengecekan caption dilewati: {exc}",
            skip_reason=str(exc),
        ))

    return issues, checks


def _check_caption_line_spacing(
    docx_path: Path,
    metadata: DocumentMetadata,
    doc: DocxDocument | None = None,
) -> tuple[list[ValidationIssue], list[ValidationCheckResult]]:
    """Validasi spasi baris pada paragraf caption gambar, tabel, dan lampiran."""
    issues: list[ValidationIssue] = []
    checks: list[ValidationCheckResult] = []

    ft = metadata.figures_and_tables
    if ft is None:
        return issues, checks

    rule_exp = ft.caption_line_spacing_rule
    val_exp  = ft.caption_line_spacing

    if rule_exp is None and val_exp is None:
        return issues, checks

    try:
        doc = doc or DocxDocument(str(docx_path))

        _CAPTION_RES = [_FIG_DETECT_RE, _TBL_DETECT_RE, _LAMPIRAN_BROAD_RE]

        caption_paras: list = [
            para for para in doc.paragraphs
            if any(rx.match(para.text.strip()) for rx in _CAPTION_RES)
        ]

        if not caption_paras:
            checks.append(ValidationCheckResult(
                category="figures_tables", field="caption_line_spacing",
                status="skipped",
                message="Tidak ada paragraf caption gambar/tabel/lampiran ditemukan",
                skip_reason="Tidak ada caption",
            ))
            return issues, checks

        def _read_line_spacing_raw(para):
            pPr = para._element.find(qn("w:pPr"))
            if pPr is None:
                return None, None
            sp_el = pPr.find(qn("w:spacing"))
            if sp_el is None:
                return None, None
            line_str  = sp_el.get(qn("w:line"))
            line_rule = sp_el.get(qn("w:lineRule"))
            if line_str is None:
                return None, None
            val = int(line_str)
            if line_rule in ("auto", None):
                mult = val / 240
                if val == 240:
                    return "SINGLE", 1.0
                if val == 360:
                    return "ONE_POINT_FIVE", 1.5
                if val == 480:
                    return "DOUBLE", 2.0
                return "MULTIPLE", mult
            if line_rule == "atLeast":
                return "AT_LEAST", val / 20
            if line_rule == "exact":
                return "EXACTLY", val / 20
            return None, None

        _RULE_LABEL: dict[str, str] = {
            "SINGLE":         "Spasi tunggal (1.0)",
            "ONE_POINT_FIVE": "Satu setengah (1.5)",
            "DOUBLE":         "Spasi ganda (2.0)",
            "MULTIPLE":       "Kelipatan",
            "AT_LEAST":       "Setidaknya",
            "EXACTLY":        "Tepat",
        }

        mismatches: list[dict] = []
        for para in caption_paras:
            rule_a, val_a = _read_line_spacing_raw(para)
            rule_ok = (rule_exp is None) or (rule_a == rule_exp.upper())
            val_ok  = (val_exp is None or val_a is None) or abs(float(val_a) - float(val_exp)) < 0.1
            if not rule_ok or not val_ok:
                mismatches.append({
                    "para_idx":  None,
                    "style":     para.style.name,
                    "text":      para.text.strip()[:100],
                    "full_text": para.text.strip(),
                    "actual":    f"{rule_a} {val_a:.2f}" if val_a else str(rule_a),
                    "bab":       None,
                    "page":      None,
                })

        rule_lbl = _RULE_LABEL.get((rule_exp or "").upper(), rule_exp or "")
        exp_str  = f"{rule_lbl} {val_exp}" if val_exp else rule_lbl

        if mismatches:
            actual_str = mismatches[0]["actual"]
            msg = (
                f"Spasi baris caption tidak sesuai (ekspektasi: {exp_str}). "
                f"{len(mismatches)} caption tidak sesuai."
            )
            occs = _build_occurrences(
                mismatches, actual_str=actual_str, expected_str=exp_str
            ) or None
            issues.append(ValidationIssue(
                category="figures_tables", field="caption_line_spacing",
                severity="error", message=msg,
                expected=exp_str, actual=actual_str,
                occurrences=occs,
            ))
            checks.append(ValidationCheckResult(
                category="figures_tables", field="caption_line_spacing",
                status="failed", message=msg,
                expected=exp_str, actual=actual_str,
                occurrences=occs,
            ))
        else:
            checks.append(ValidationCheckResult(
                category="figures_tables", field="caption_line_spacing",
                status="passed",
                message=f"Spasi baris caption ({exp_str}): semua sesuai",
                expected=exp_str, actual=exp_str,
            ))

    except Exception as exc:
        checks.append(ValidationCheckResult(
            category="figures_tables", field="caption_line_spacing",
            status="skipped",
            message=f"Pengecekan spasi caption dilewati: {exc}",
            skip_reason=str(exc),
        ))

    return issues, checks


def _check_budget_format(
    docx_path: Path,
    metadata: DocumentMetadata,
    doc: DocxDocument | None = None,
) -> tuple[list[ValidationIssue], list[ValidationCheckResult]]:
    """Validasi keberadaan dan kategori tabel anggaran (RAB) pada dokumen proposal.

    Mengecek apakah tabel anggaran ditemukan dan apakah kategori pengeluaran
    yang ditetapkan metadata muncul di dalamnya. Tidak memvalidasi nilai nominal.
    """
    issues: list[ValidationIssue] = []
    checks: list[ValidationCheckResult] = []

    ft = metadata.figures_and_tables
    if ft is None or ft.budget_format_rules is None:
        return issues, checks

    budget_rules = ft.budget_format_rules
    budget_items = budget_rules.budget_items or []

    try:
        doc = doc or DocxDocument(str(docx_path))

        _BUDGET_KEYWORDS = {
            "jenis pengeluaran", "biaya", "anggaran", "rab",
            "pengeluaran", "peralatan", "bahan", "perjalanan", "sewa",
        }

        budget_table = None
        for tbl in doc.tables:
            if not tbl.rows:
                continue
            header_text = " ".join(
                c.text.strip().lower() for c in tbl.rows[0].cells
            )
            if any(kw in header_text for kw in _BUDGET_KEYWORDS):
                budget_table = tbl
                break

        if budget_table is None:
            msg = "Tabel anggaran (RAB) tidak ditemukan di dokumen."
            issues.append(ValidationIssue(
                category="figures_tables", field="budget_format",
                severity="warning", message=msg,
                expected="tabel anggaran", actual="tidak ditemukan",
            ))
            checks.append(ValidationCheckResult(
                category="figures_tables", field="budget_format",
                status="failed", message=msg,
                expected="tabel anggaran", actual="tidak ditemukan",
            ))
            return issues, checks

        if not budget_items:
            checks.append(ValidationCheckResult(
                category="figures_tables", field="budget_format",
                status="passed",
                message="Tabel anggaran ditemukan di dokumen.",
                expected="tabel anggaran", actual="ditemukan",
            ))
            return issues, checks

        all_table_text = " ".join(
            c.text.strip().lower()
            for row in budget_table.rows
            for c in row.cells
        )

        missing_categories: list[str] = []
        for item in budget_items:
            jenis = (item.jenis_pengeluaran or "").strip().lower()
            if jenis and jenis not in all_table_text:
                missing_categories.append(item.jenis_pengeluaran or "")

        if missing_categories:
            missing_str = ", ".join(missing_categories[:5])
            msg = (
                f"Kategori anggaran berikut tidak ditemukan di tabel RAB: {missing_str}."
            )
            issues.append(ValidationIssue(
                category="figures_tables", field="budget_format",
                severity="warning", message=msg,
                expected=", ".join(
                    i.jenis_pengeluaran or "" for i in budget_items
                ),
                actual=f"tidak ada: {missing_str}",
            ))
            checks.append(ValidationCheckResult(
                category="figures_tables", field="budget_format",
                status="failed", message=msg,
                expected=", ".join(i.jenis_pengeluaran or "" for i in budget_items),
                actual=f"tidak ada: {missing_str}",
            ))
        else:
            checks.append(ValidationCheckResult(
                category="figures_tables", field="budget_format",
                status="passed",
                message="Tabel anggaran ditemukan dan semua kategori pengeluaran ada.",
                expected=", ".join(i.jenis_pengeluaran or "" for i in budget_items),
                actual="semua kategori ditemukan",
            ))

    except Exception as exc:
        checks.append(ValidationCheckResult(
            category="figures_tables", field="budget_format",
            status="skipped",
            message=f"Pengecekan tabel anggaran dilewati: {exc}",
            skip_reason=str(exc),
        ))

    return issues, checks
