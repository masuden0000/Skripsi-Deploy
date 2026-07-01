"""Structure checks: document structure order and lampiran format. Keyword: automated document validation"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

from model_ai.extractor.models import DocumentMetadata
from model_ai.validation.models import ValidationCheckResult, ValidationIssue
from model_ai.validation.validocx_adapter import (
    _resolve_line_spacing,
    _heading_level_from_style,
)
from ._shared import (
    _CAPTION_ALIGN_MAP,
    _BAB_RE,
    _SUB_BAB_RE,
    _LAMPIRAN_ITEM_RE,
    _LAMPIRAN_BROAD_RE,
    _TOC_TOF_STYLE_NAMES,
    _HEADING_TITLE_MAP,
    _HEADING_TITLE_MAP_INV,
    _build_lampiran_re,
    _build_occurrences,
)


def _classify_heading(
    text: str,
    lampiran_re: re.Pattern | None = None,
) -> tuple[str | None, dict]:
    """Klasifikasi teks heading menjadi tipe section + info tambahan.

    lampiran_re: regex dinamis dari metadata (separator per skema).
                 Bila None, fallback ke _LAMPIRAN_ITEM_RE (titik).
    """
    text_stripped = text.strip()
    text_upper = text_stripped.upper()

    if text_upper in _HEADING_TITLE_MAP:
        return _HEADING_TITLE_MAP[text_upper], {}

    m = _BAB_RE.match(text_upper)
    if m:
        return "bab", {"number": int(m.group(1))}

    m = _SUB_BAB_RE.match(text_stripped)
    if m:
        return "sub_bab", {"sub_number": f"{m.group(1)}.{m.group(2)}"}

    m = (lampiran_re or _LAMPIRAN_ITEM_RE).match(text_stripped)
    if m:
        return "item_lampiran", {"lampiran_number": f"Lampiran {m.group(1)}"}

    return None, {}


def _detect_artikel_pre_bab_sections(doc: DocxDocument, actual_classified: list[dict]) -> None:
    """Deteksi section halaman pertama artikel PKM-AI yang tidak pakai heading style.

    Layer 1 — position-based: jika ada paragraf non-kosong sebelum BAB 1 maka
    judul dan identitas_penulis dianggap ada (pola tetap 2023–2026).
    Layer 2 — plain-text fallback: scan teks sebelum BAB 1 untuk ABSTRAK/ABSTRACT
    jika belum terdeteksi via heading map (layer 1).
    """
    already_types = {s["type"] for s in actual_classified}

    bab1_idx: int | None = None
    for i, para in enumerate(doc.paragraphs):
        if _BAB_RE.match(para.text.strip().upper()):
            bab1_idx = i
            break

    if bab1_idx is None:
        return

    pre_bab_paras = [p for p in doc.paragraphs[:bab1_idx] if p.text.strip()]
    if not pre_bab_paras:
        return

    if "judul" not in already_types:
        actual_classified.insert(0, {"type": "judul", "text": pre_bab_paras[0].text.strip()})
        already_types.add("judul")
    if "identitas_penulis" not in already_types:
        actual_classified.insert(1, {"type": "identitas_penulis", "text": "(identitas penulis)"})
        already_types.add("identitas_penulis")

    for para in pre_bab_paras:
        text_upper = para.text.strip().upper()
        if "abstrak" not in already_types and text_upper == "ABSTRAK":
            actual_classified.append({"type": "abstrak", "text": para.text.strip()})
            already_types.add("abstrak")
        if "abstract" not in already_types and text_upper == "ABSTRACT":
            actual_classified.append({"type": "abstract", "text": para.text.strip()})
            already_types.add("abstract")


def _check_document_structure(
    docx_path: Path,
    metadata: DocumentMetadata,
    doc: DocxDocument | None = None,
) -> tuple[list[ValidationIssue], list[ValidationCheckResult]]:
    """Validasi urutan dan kehadiran section dokumen berdasarkan document_structure_proposal."""
    issues: list[ValidationIssue] = []
    checks: list[ValidationCheckResult] = []

    _ds_p = metadata.document_structure_proposal
    _ds_a = metadata.document_structure_artikel
    if _ds_p and _ds_p.sections:
        ds = _ds_p
    elif _ds_a and _ds_a.sections:
        ds = _ds_a
    else:
        ds = _ds_p or _ds_a
    if ds is None or not ds.sections:
        checks.append(ValidationCheckResult(
            category="document_structure", field="section_order",
            status="skipped",
            message="Tidak ada data document_structure di metadata",
            skip_reason="Tidak ada nilai di metadata",
        ))
        return issues, checks

    try:
        doc = doc or DocxDocument(str(docx_path))

        _ds_sep = ds.lampiran_heading_separator if ds else None
        _lampiran_re = _build_lampiran_re(_ds_sep if _ds_sep is not None else ".")

        actual_classified: list[dict] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            level = _heading_level_from_style(para.style)
            if level is None:
                continue
            section_type, extra = _classify_heading(text, lampiran_re=_lampiran_re)
            if section_type:
                actual_classified.append({"type": section_type, "text": text, **extra})

        if ds is _ds_a:
            _detect_artikel_pre_bab_sections(doc, actual_classified)

        expected_major = [s for s in ds.sections if s.is_major_section]
        required_non_bab_types = {
            s.type for s in expected_major if s.required is not False and s.type != "bab"
        }
        required_bab_nums = {
            s.number for s in expected_major
            if s.type == "bab" and s.required is not False and s.number is not None
        }
        actual_types_set = {s["type"] for s in actual_classified}
        actual_bab_nums  = {
            s["number"] for s in actual_classified
            if s["type"] == "bab" and "number" in s
        }

        def _req_label(section_type: str, title: str | None) -> str:
            """Label terbaca manusia untuk sebuah tipe section."""
            if section_type == "bab":
                return "BAB"
            if section_type == "sub_bab":
                return "Sub BAB"
            if title:
                return title
            return _HEADING_TITLE_MAP_INV.get(section_type, section_type.replace("_", " ").upper())

        seen_meta: set[str] = set()
        required_labels: list[str] = []
        optional_labels: list[str] = []
        for s in expected_major:
            if s.type in seen_meta:
                continue
            seen_meta.add(s.type)
            lbl = _req_label(s.type, s.title)
            if s.required is not False:
                required_labels.append(lbl)
            else:
                optional_labels.append(lbl)

        major_types_set = {s.type for s in expected_major}
        seen_actual: set[str] = set()
        actual_section_texts: list[str] = []
        for s in actual_classified:
            if s["type"] not in major_types_set:
                continue
            if s["type"] == "bab":
                actual_section_texts.append(s["text"])
            elif s["type"] not in seen_actual:
                actual_section_texts.append(s["text"])
                seen_actual.add(s["type"])

        req_part = ', '.join(required_labels) if required_labels else "–"
        opt_part = ', '.join(optional_labels) if optional_labels else "–"
        expected_display_req = f"Wajib: {req_part} | Opsional: {opt_part}"
        actual_display_req   = ', '.join(actual_section_texts) if actual_section_texts else "Tidak ada"

        actual_major_found = [s for s in actual_classified if s["type"] in major_types_set]
        occ_req = _build_occurrences(
            [{"text": s["text"][:100], "full_text": s["text"],
              "style": "", "page": None, "bab": None, "para_idx": None}
             for s in actual_major_found],
            actual_str=None, expected_str=None,
        ) or None

        missing_non_bab = required_non_bab_types - actual_types_set
        missing_bab_nums = required_bab_nums - actual_bab_nums
        total_required   = len(required_non_bab_types) + len(required_bab_nums)

        if missing_non_bab or missing_bab_nums:
            missing_labels: list[str] = []
            seen_ml: set[str] = set()
            for s in expected_major:
                if s.type != "bab" and s.type in missing_non_bab and s.type not in seen_ml:
                    seen_ml.add(s.type)
                    missing_labels.append(_req_label(s.type, s.title))
            for s in sorted(
                [x for x in expected_major if x.type == "bab" and x.number in missing_bab_nums],
                key=lambda x: x.number or 0,
            ):
                lbl = f"BAB {s.number} {s.title.upper()}" if s.title else f"BAB {s.number}"
                missing_labels.append(lbl)

            total_missing = len(missing_non_bab) + len(missing_bab_nums)
            msg = f"{total_missing} section wajib tidak ditemukan: {', '.join(missing_labels)}"
            issues.append(ValidationIssue(
                category="document_structure", field="required_section",
                severity="error", message=msg,
                expected=expected_display_req, actual=actual_display_req,
            ))
            checks.append(ValidationCheckResult(
                category="document_structure", field="required_section",
                status="failed", message=msg,
                expected=expected_display_req, actual=actual_display_req,
                occurrences=occ_req,
            ))
        else:
            checks.append(ValidationCheckResult(
                category="document_structure", field="required_section",
                status="passed",
                message=f"Semua {total_required} section wajib ditemukan",
                expected=expected_display_req,
                actual=actual_display_req,
                occurrences=occ_req,
            ))

        bab_actuals = [s for s in actual_classified if s["type"] == "bab"]
        bab_numbers = [s["number"] for s in bab_actuals if "number" in s]
        expected_bab_sections = sorted(
            [s for s in expected_major if s.type == "bab" and s.number is not None],
            key=lambda s: s.number,
        )
        expected_bab_numbers = [s.number for s in expected_bab_sections]

        def _bab_label(number: int, title: str | None) -> str:
            """Format label BAB: 'BAB N JUDUL' atau fallback 'BAB N'."""
            if title:
                return f"BAB {number} {title.upper()}"
            return f"BAB {number}"

        expected_bab_labels = [_bab_label(s.number, s.title) for s in expected_bab_sections]
        actual_bab_labels   = [s["text"] for s in bab_actuals]
        bab_num_to_text: dict[int, str] = {
            s["number"]: s["text"] for s in bab_actuals if "number" in s
        }

        if expected_bab_numbers and bab_numbers:
            if bab_numbers != sorted(bab_numbers):
                actual_ordered_labels   = [bab_num_to_text.get(n, f"BAB {n}") for n in bab_numbers]
                expected_sorted_labels  = [_bab_label(s.number, s.title) for s in expected_bab_sections]
                msg = f"BAB tidak berurutan. Ditemukan: {' → '.join(actual_ordered_labels)}"
                occ_bab_err = _build_occurrences(
                    [{"text": s["text"][:100], "full_text": s["text"],
                      "style": "", "page": None, "bab": None, "para_idx": None}
                     for s in bab_actuals],
                    actual_str=None, expected_str=None,
                ) or None
                issues.append(ValidationIssue(
                    category="document_structure", field="bab_order",
                    severity="error", message=msg,
                    expected=' → '.join(expected_sorted_labels),
                    actual=' → '.join(actual_ordered_labels),
                ))
                checks.append(ValidationCheckResult(
                    category="document_structure", field="bab_order",
                    status="failed", message=msg,
                    expected=' → '.join(expected_sorted_labels),
                    actual=' → '.join(actual_ordered_labels),
                    occurrences=occ_bab_err,
                ))
            else:
                missing_bab_nums = set(expected_bab_numbers) - set(bab_numbers)
                if missing_bab_nums:
                    missing_labels = [
                        _bab_label(s.number, s.title)
                        for s in expected_bab_sections if s.number in missing_bab_nums
                    ]
                    msg = f"BAB berikut tidak ditemukan: {', '.join(missing_labels)}"
                    occ_bab_missing = _build_occurrences(
                        [{"text": s["text"][:100], "full_text": s["text"],
                          "style": "", "page": None, "bab": None, "para_idx": None}
                         for s in bab_actuals],
                        actual_str=None, expected_str=None,
                    ) or None
                    issues.append(ValidationIssue(
                        category="document_structure", field="bab_order",
                        severity="error", message=msg,
                        expected=' → '.join(expected_bab_labels),
                        actual=' → '.join(actual_bab_labels),
                    ))
                    checks.append(ValidationCheckResult(
                        category="document_structure", field="bab_order",
                        status="failed", message=msg,
                        expected=' → '.join(expected_bab_labels),
                        actual=' → '.join(actual_bab_labels),
                        occurrences=occ_bab_missing,
                    ))
                else:
                    # Cek judul BAB: nomor urut benar, tapi judul harus sesuai skema
                    _bab_title_re = re.compile(r'^BAB\s+\d+[\s.]*', re.IGNORECASE)
                    bab_title_mismatches: list[tuple] = []
                    for _exp_bab in expected_bab_sections:
                        _actual_text = bab_num_to_text.get(_exp_bab.number, "")
                        if _exp_bab.title and _actual_text:
                            _actual_title = _bab_title_re.sub("", _actual_text).strip().upper()
                            _expected_title = _exp_bab.title.strip().upper()
                            if _actual_title != _expected_title:
                                bab_title_mismatches.append(
                                    (_exp_bab.number, _exp_bab.title, _actual_text)
                                )
                    occ_bab = _build_occurrences(
                        [{"text": s["text"][:100], "full_text": s["text"],
                          "style": "", "page": None, "bab": None, "para_idx": None}
                         for s in bab_actuals],
                        actual_str=None, expected_str=None,
                    ) or None
                    if bab_title_mismatches:
                        _wrong_parts = [
                            f"BAB {num}: ditemukan '{act}', seharusnya '{_bab_label(num, exp)}'"
                            for num, exp, act in bab_title_mismatches
                        ]
                        msg = f"Judul BAB tidak sesuai skema. {'; '.join(_wrong_parts)}"
                        issues.append(ValidationIssue(
                            category="document_structure", field="bab_order",
                            severity="error", message=msg,
                            expected=' → '.join(expected_bab_labels),
                            actual=' → '.join(actual_bab_labels),
                        ))
                        checks.append(ValidationCheckResult(
                            category="document_structure", field="bab_order",
                            status="failed", message=msg,
                            expected=' → '.join(expected_bab_labels),
                            actual=' → '.join(actual_bab_labels),
                            occurrences=occ_bab,
                        ))
                    else:
                        checks.append(ValidationCheckResult(
                            category="document_structure", field="bab_order",
                            status="passed",
                            message=f"BAB berurutan dengan benar: {' → '.join(actual_bab_labels)}",
                            expected=' → '.join(expected_bab_labels),
                            actual=' → '.join(actual_bab_labels),
                            occurrences=occ_bab,
                        ))

        def _section_label(section_type: str, title: str | None = None) -> str:
            if section_type == "sub_bab":
                return "Sub BAB"
            if title:
                return title
            return _HEADING_TITLE_MAP_INV.get(section_type) or section_type.replace("_", " ").upper()

        type_to_label: dict[str, str] = {}
        for s in expected_major:
            if s.type not in type_to_label:
                type_to_label[s.type] = _section_label(s.type, s.title)

        def _expand_labels(type_list: list[str], bab_labels: list[str]) -> str:
            parts: list[str] = []
            for t in type_list:
                if t == "bab":
                    parts.extend(bab_labels if bab_labels else ["BAB"])
                else:
                    parts.append(type_to_label.get(t, t.replace("_", " ").upper()))
            return ' → '.join(parts)

        seen: set[str] = set()
        actual_order: list[str] = []
        for s in actual_classified:
            key = s["type"]
            if key not in seen:
                actual_order.append(key)
                seen.add(key)

        seen = set()
        expected_order: list[str] = []
        for s in expected_major:
            key = s.type
            if key not in seen:
                expected_order.append(key)
                seen.add(key)

        expected_filtered = [t for t in expected_order if t in actual_types_set]
        actual_filtered = [t for t in actual_order if t in set(expected_order)]

        exp_display = _expand_labels(expected_filtered, expected_bab_labels)
        act_display = _expand_labels(actual_filtered, actual_bab_labels)

        if expected_filtered and actual_filtered and actual_filtered != expected_filtered:
            msg = (
                f"Urutan section tidak sesuai. "
                f"Seharusnya: {exp_display}, "
                f"Ditemukan: {act_display}"
            )
            occ_sec_err = _build_occurrences(
                [{"text": s["text"][:100], "full_text": s["text"],
                  "style": "", "page": None, "bab": None, "para_idx": None}
                 for s in actual_classified if s["type"] in set(expected_order)],
                actual_str=None, expected_str=None,
            ) or None
            issues.append(ValidationIssue(
                category="document_structure", field="section_order",
                severity="error", message=msg,
                expected=exp_display,
                actual=act_display,
            ))
            checks.append(ValidationCheckResult(
                category="document_structure", field="section_order",
                status="failed", message=msg,
                expected=exp_display,
                actual=act_display,
                occurrences=occ_sec_err,
            ))
        elif expected_filtered:
            major_found = [s for s in actual_classified if s["type"] in set(expected_order)]
            occ_sec = _build_occurrences(
                [{"text": s["text"][:100], "full_text": s["text"],
                  "style": "", "page": None, "bab": None, "para_idx": None}
                 for s in major_found],
                actual_str=None, expected_str=None,
            ) or None
            checks.append(ValidationCheckResult(
                category="document_structure", field="section_order",
                status="passed",
                message=f"Urutan section sesuai: {act_display}",
                expected=exp_display,
                actual=act_display,
                occurrences=occ_sec,
            ))

    except Exception as exc:
        checks.append(ValidationCheckResult(
            category="document_structure", field="section_order",
            status="skipped",
            message=f"Pengecekan struktur dokumen dilewati: {exc}",
            skip_reason=str(exc),
        ))

    return issues, checks


def _check_lampiran_format(
    docx_path: Path,
    metadata: DocumentMetadata,
    doc: DocxDocument | None = None,
) -> tuple[list[ValidationIssue], list[ValidationCheckResult]]:
    """Validasi judul lampiran via text-pattern untuk SEMUA judul lampiran.

    Mendeteksi paragraf yang teksnya diawali 'Lampiran <angka>' (_LAMPIRAN_BROAD_RE),
    baik yang menggunakan style bernama 'Lampiran' maupun style lain (Normal, Body, dll).
    Aturan atribut: font family, font size, line spacing, alignment JUSTIFY — sama dengan body.
    """
    issues: list[ValidationIssue] = []
    checks: list[ValidationCheckResult] = []

    t  = metadata.typography
    _ds_p2 = metadata.document_structure_proposal
    _ds_a2 = metadata.document_structure_artikel
    if _ds_p2 and _ds_p2.sections:
        ds = _ds_p2
    elif _ds_a2 and _ds_a2.sections:
        ds = _ds_a2
    else:
        ds = _ds_p2 or _ds_a2
    expected_font    = t.font_family if t else None
    expected_size    = int(t.font_size_body_pt) if t and t.font_size_body_pt else None
    expected_spacing = _resolve_line_spacing(metadata)

    _body_align_str_lmp = (metadata.spacing.paragraph_alignment if metadata.spacing else None) or "JUSTIFY"
    _expected_align_lmp = _CAPTION_ALIGN_MAP.get(_body_align_str_lmp.upper(), WD_ALIGN_PARAGRAPH.JUSTIFY)
    separator   = ds.lampiran_heading_separator if ds else None
    effective_sep = separator if separator is not None else "."
    lampiran_re = _build_lampiran_re(effective_sep)

    expected_summary = ", ".join(filter(None, [
        expected_font or None,
        f"{expected_size}pt" if expected_size else None,
        f"spacing {expected_spacing}",
        _body_align_str_lmp,
    ]))

    try:
        doc = doc or DocxDocument(str(docx_path))

        pass_items:      list[dict] = []
        sep_pass_items:  list[dict] = []
        wrong_alignment: list[dict] = []
        wrong_font:      list[dict] = []
        wrong_size:      list[dict] = []
        wrong_spacing:   list[dict] = []
        wrong_separator: list[str]  = []
        total = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text or not _LAMPIRAN_BROAD_RE.match(text):
                continue

            if para.style.name in _TOC_TOF_STYLE_NAMES:
                continue

            _sep_ok = bool(lampiran_re.match(text))
            if not _sep_ok:
                wrong_separator.append(text[:70])

            total += 1
            para_info: dict = {
                "text"      : text[:100],
                "full_text" : text,
                "style"     : para.style.name,
                "page"      : None,
                "bab"       : None,
                "para_idx"  : None,
            }
            has_issue = False
            if _sep_ok:
                sep_pass_items.append(para_info)

            align = para.paragraph_format.alignment
            if align is None:
                try:
                    align = para.style.paragraph_format.alignment
                except Exception:
                    align = None
            if align is not None and align != _expected_align_lmp:
                _align_names = {0: "LEFT", 1: "CENTER", 2: "RIGHT", 3: "JUSTIFY"}
                wrong_alignment.append({**para_info, "actual": _align_names.get(int(align), str(align))})
                has_issue = True

            for run in para.runs:
                if expected_font and run.font.name and run.font.name != expected_font:
                    wrong_font.append({**para_info, "actual": run.font.name})
                    has_issue = True
                    break
                if expected_size and run.font.size:
                    actual_pt = round(run.font.size.pt)
                    if actual_pt != expected_size:
                        wrong_size.append({**para_info, "actual": f"{actual_pt}pt"})
                        has_issue = True
                        break

            ls = para.paragraph_format.line_spacing
            if ls is not None:
                try:
                    ls_val = round(float(ls), 2)
                    if abs(ls_val - expected_spacing) > 0.05:
                        wrong_spacing.append({**para_info, "actual": str(ls_val)})
                        has_issue = True
                except (TypeError, ValueError):
                    pass

            if not has_issue:
                pass_items.append(para_info)

        sep_display = f'titik (".")' if effective_sep == "." else (
            f'"{effective_sep}"' if effective_sep else "tanpa titik"
        )
        if wrong_separator:
            msg = (
                f"{len(wrong_separator)} judul lampiran tidak menggunakan format yang diharapkan "
                f"({sep_display} setelah nomor). Contoh: \"{wrong_separator[0]}\""
            )
            occ_sep_err = _build_occurrences(
                [{"text": t[:100], "full_text": t, "style": "",
                  "page": None, "bab": None, "para_idx": None}
                 for t in wrong_separator],
                actual_str=None, expected_str=effective_sep,
            ) or None
            issues.append(ValidationIssue(
                category="document_structure", field="lampiran_separator",
                severity="error", message=msg,
                expected=effective_sep, actual=wrong_separator[0],
            ))
            checks.append(ValidationCheckResult(
                category="document_structure", field="lampiran_separator",
                status="failed", message=msg,
                expected=effective_sep, actual=wrong_separator[0],
                occurrences=occ_sep_err,
            ))
        else:
            occ_sep = _build_occurrences(
                sep_pass_items, actual_str=effective_sep, expected_str=effective_sep,
            ) or None
            checks.append(ValidationCheckResult(
                category="document_structure", field="lampiran_separator",
                status="passed",
                message=f"Format penulisan judul lampiran sesuai ({sep_display} setelah nomor)",
                expected=effective_sep,
                occurrences=occ_sep,
            ))

        if total == 0:
            checks.append(ValidationCheckResult(
                category="typography", field="lampiran_format",
                status="skipped",
                message="Tidak ada judul lampiran yang ditemukan di dokumen",
                skip_reason="Tidak ada paragraf dengan pola 'Lampiran <angka>'",
            ))
            return issues, checks

        all_ok = not any([wrong_alignment, wrong_font, wrong_size, wrong_spacing])
        if pass_items:
            n_pass = len(pass_items)
            occs_pass = _build_occurrences(
                pass_items,
                actual_str=expected_summary,
                expected_str=expected_summary,
            ) or None
            checks.append(ValidationCheckResult(
                category="typography", field="lampiran_format",
                status="passed",
                message=(
                    f"Semua {total} judul lampiran sesuai format body"
                    if all_ok else
                    f"{n_pass} dari {total} judul lampiran sesuai format body"
                ),
                expected=expected_summary,
                actual=expected_summary,
                occurrences=occs_pass,
            ))

        for field, items, label, expected_val, is_error in [
            ("lampiran_alignment", wrong_alignment, "alignment",    _body_align_str_lmp,    True),
            ("lampiran_font",      wrong_font,      "font family",  expected_font or "",    False),
            ("lampiran_font_size", wrong_size,      "font size",    f"{expected_size}pt",   False),
            ("lampiran_spacing",   wrong_spacing,   "line spacing", str(expected_spacing),  False),
        ]:
            if not items:
                continue
            first_actual = items[0].get("actual", "")
            msg = (
                f"{len(items)} judul lampiran {label} tidak sesuai "
                f"(ekspektasi: {expected_val}). Contoh: \"{items[0]['text']}\""
            )
            occs_fail = _build_occurrences(items, expected_str=str(expected_val)) or None
            issues.append(ValidationIssue(
                category="typography", field=field,
                severity="error" if is_error else "warning", message=msg,
                expected=str(expected_val), actual=first_actual,
                occurrences=occs_fail,
            ))
            checks.append(ValidationCheckResult(
                category="typography", field=field,
                status="failed" if is_error else "warning", message=msg,
                expected=str(expected_val), actual=first_actual,
                occurrences=occs_fail,
            ))

    except Exception as exc:
        checks.append(ValidationCheckResult(
            category="typography", field="lampiran_format",
            status="skipped",
            message=f"Pengecekan format lampiran dilewati: {exc}",
            skip_reason=str(exc),
        ))

    return issues, checks
