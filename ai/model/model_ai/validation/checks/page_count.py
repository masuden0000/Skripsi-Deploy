"""Page count checks: structural page counting and section start validation. Keyword: automated document validation"""
from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from model_ai.extractor.models import DocumentMetadata
from model_ai.validation.models import ValidationCheckResult, ValidationIssue
from model_ai.validation.validocx_adapter import _heading_level_from_style

from ._shared import (
    _BAB_RE,
    _UNREALISTIC_PAGE_FACTOR,
    _HEADING_TITLE_MAP_INV,
    _build_occurrences,
)
from .numbering import _read_section_page_numbering


def _count_pages_structural(doc) -> dict[int, int]:
    """Peta para_index → page_number (1-based) menggunakan penanda page break di XML.

    Strategi (dua pass):
      Pass 1 (utama) : w:lastRenderedPageBreak — disimpan Word setelah rendering,
                       paling andal untuk dokumen yang sudah disimpan oleh Word.
      Pass 2 (fallback): jika pass 1 menghasilkan max_page == 1 (tidak ada
                        rendered page break), gunakan explicit w:br type="page"
                        + inline w:sectPr saja.

    Inline w:sectPr di w:pPr suatu paragraf menandai section break:
    paragraf itu sendiri tetap di halaman yang sama, tapi paragraf BERIKUTNYA
    mulai di halaman baru.
    """
    para_list = list(doc.paragraphs)
    if not para_list:
        return {}

    def _build_structural_map(prefer_rendered_page_breaks: bool) -> dict[int, int]:
        result: dict[int, int] = {}
        current_page = 1
        for idx, para in enumerate(para_list):
            p = para._p
            has_page_break = False

            if prefer_rendered_page_breaks and idx > 0:
                if p.findall(".//" + qn("w:lastRenderedPageBreak")):
                    has_page_break = True

            if not has_page_break:
                for br in p.findall(".//" + qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        has_page_break = True
                        break

            if has_page_break and idx > 0:
                current_page += 1

            result[idx] = current_page

            pPr = p.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                current_page += 1

        return result

    page_map = _build_structural_map(prefer_rendered_page_breaks=True)
    if max(page_map.values(), default=1) <= 1:
        page_map = _build_structural_map(prefer_rendered_page_breaks=False)
    return page_map


def _build_displayed_page_map(doc) -> dict[int, int]:
    """Peta para_index → nomor halaman yang ditampilkan di header/footer.

    Membaca w:pgNumType w:start dari tiap section sehingga nomor halaman
    yang dihasilkan sesuai dengan yang tertera di header dokumen:
      - Section preliminary (romawi): i, ii, iii, ...
      - Section konten (arab): 1, 2, 3, ...

    Deteksi page break: paragraf dianggap memulai halaman baru bila punya
    w:lastRenderedPageBreak (rendered page break, disimpan Word otomatis saat
    render) ATAU w:br type="page" (explicit page break, dipasang manual oleh
    user). Keduanya diperlakukan setara per paragraf — audit menunjukkan tidak
    ada paragraf yang punya keduanya sekaligus, sehingga tidak ada risiko
    menghitung transisi halaman yang sama dua kali.

    Iterasi mencakup SELURUH paragraf termasuk yang ada di dalam sel tabel,
    karena Word menyimpan w:lastRenderedPageBreak di dalam sel tabel ketika
    tabel tersebut membentang antar halaman.  Hanya paragraf body-level
    (doc.paragraphs) yang dimasukkan ke dalam result.
    """
    para_list = list(doc.paragraphs)
    if not para_list:
        return {}

    body_para_ids: dict[int, int] = {id(p._p): i for i, p in enumerate(para_list)}

    section_ends: list[tuple[int, object]] = []
    for idx, para in enumerate(para_list):
        pPr = para._p.find(qn("w:pPr"))
        if pPr is not None:
            sectPr = pPr.find(qn("w:sectPr"))
            if sectPr is not None:
                section_ends.append((idx, sectPr))

    try:
        body_sectPr = doc.element.body.find(qn("w:sectPr"))
    except Exception:
        body_sectPr = None
    section_ends.append((len(para_list) - 1, body_sectPr))

    result: dict[int, int] = {}
    sec_start = 0

    for end_idx, sectPr in section_ends:
        if end_idx < sec_start:
            continue

        if sectPr is not None:
            _, section_start_page = _read_section_page_numbering(sectPr)
        else:
            section_start_page = 1

        current_page = section_start_page

        start_elem = para_list[sec_start]._p
        end_elem   = para_list[end_idx]._p

        found_start     = False
        first_in_section = True

        for p_elem in doc.element.body.iter(qn("w:p")):
            if not found_start:
                if p_elem is start_elem:
                    found_start = True
                else:
                    continue

            if not first_in_section:
                has_break = bool(
                    p_elem.findall(".//" + qn("w:lastRenderedPageBreak"))
                ) or any(
                    br.get(qn("w:type")) == "page"
                    for br in p_elem.findall(".//" + qn("w:br"))
                )
                if has_break:
                    current_page += 1
            else:
                first_in_section = False

            body_idx = body_para_ids.get(id(p_elem))
            if body_idx is not None:
                result[body_idx] = current_page

            if p_elem is end_elem:
                break

        sec_start = end_idx + 1

    return result


def _find_section_para_idx(
    para_list: list,
    section_type: str,
    search_from: int = 0,
) -> tuple[int | None, str]:
    """Cari indeks paragraf heading pertama yang cocok dengan section_type.

    Mengembalikan (para_idx, text) atau (None, "") jika tidak ditemukan.
    Hanya paragraf dengan style heading (deteksi via _heading_level_from_style)
    yang diperiksa — entri TOC/TOF (style "toc 1", Normal, dsb.) diabaikan.
    """
    if section_type == "bab":
        for i in range(search_from, len(para_list)):
            para = para_list[i]
            text = para.text.strip()
            if not text:
                continue
            if _heading_level_from_style(para.style) is None:
                continue
            if _BAB_RE.match(text.upper()):
                return i, text
    else:
        expected_title = _HEADING_TITLE_MAP_INV.get(
            section_type, section_type.upper().replace("_", " ")
        )
        for i in range(search_from, len(para_list)):
            para = para_list[i]
            text = para.text.strip()
            if not text:
                continue
            if _heading_level_from_style(para.style) is None:
                continue
            if text.upper() == expected_title:
                return i, text
    return None, ""


def _check_page_count(
    docx_path: Path,
    metadata: DocumentMetadata,
    doc: DocxDocument | None = None,
) -> tuple[list[ValidationIssue], list[ValidationCheckResult]]:
    """Validasi jumlah halaman inti tidak melebihi batas maksimum.

    Halaman inti dihitung dari section halaman_inti_mulai (default: bab)
    hingga halaman halaman_inti_selesai (default: daftar_pustaka), INKLUSIF.
    Penghitungan halaman menggunakan penanda struktural di XML:
    w:lastRenderedPageBreak (utama), explicit page break, dan inline sectPr.
    """
    issues: list[ValidationIssue] = []
    checks: list[ValidationCheckResult] = []

    pc = metadata.page_count_limits
    if pc is None:
        checks.append(ValidationCheckResult(
            category="page_count", field="halaman_inti",
            status="skipped",
            message="Batas halaman inti tidak dikonfigurasi di metadata",
            skip_reason="page_count_limits tidak ada",
        ))
        return issues, checks

    is_artikel = pc.artikel_halaman_inti_maks is not None
    maks       = pc.artikel_halaman_inti_maks if is_artikel else pc.proposal_halaman_inti_maks
    min_pages  = pc.artikel_halaman_inti_min if is_artikel else None

    if maks is None:
        checks.append(ValidationCheckResult(
            category="page_count", field="halaman_inti",
            status="skipped",
            message="Batas maksimum halaman inti tidak dikonfigurasi di metadata",
            skip_reason="halaman_inti_maks tidak ada",
        ))
        return issues, checks

    mulai_type   = pc.halaman_inti_mulai or "bab"
    selesai_type = pc.halaman_inti_selesai or "daftar_pustaka"

    try:
        doc = doc or DocxDocument(str(docx_path))
        para_list = list(doc.paragraphs)

        start_idx, start_text = _find_section_para_idx(para_list, mulai_type)

        if start_idx is None:
            mulai_label = (
                "BAB pertama" if mulai_type == "bab"
                else _HEADING_TITLE_MAP_INV.get(mulai_type, mulai_type.upper())
            )
            checks.append(ValidationCheckResult(
                category="page_count", field="halaman_inti",
                status="skipped",
                message=(
                    f"Penghitungan halaman dilewati: heading '{mulai_label}' "
                    f"(halaman_inti_mulai='{mulai_type}') tidak ditemukan di dokumen"
                ),
                skip_reason=f"heading '{mulai_label}' tidak ditemukan",
            ))
            return issues, checks

        end_idx, end_text = _find_section_para_idx(
            para_list, selesai_type, search_from=start_idx + 1
        )

        if end_idx is None:
            selesai_label = _HEADING_TITLE_MAP_INV.get(
                selesai_type, selesai_type.upper().replace("_", " ")
            )
            checks.append(ValidationCheckResult(
                category="page_count", field="halaman_inti",
                status="skipped",
                message=(
                    f"Penghitungan halaman dilewati: heading '{selesai_label}' "
                    f"(halaman_inti_selesai='{selesai_type}') tidak ditemukan "
                    f"setelah '{start_text}'"
                ),
                skip_reason=f"heading '{selesai_label}' tidak ditemukan setelah start",
            ))
            return issues, checks

        page_map   = _build_displayed_page_map(doc)
        start_page = page_map.get(start_idx, 1)
        end_page   = page_map.get(end_idx, 1)
        count      = end_page - start_page + 1

        if count <= 0:
            checks.append(ValidationCheckResult(
                category="page_count", field="halaman_inti",
                status="skipped",
                message=(
                    f"Penghitungan halaman menghasilkan nilai tidak valid ({count}). "
                    "Urutan section mungkin tidak linear atau dokumen rusak."
                ),
                skip_reason=f"hitung halaman tidak valid: {count}",
            ))
            return issues, checks

        unrealistic_limit = maks * _UNREALISTIC_PAGE_FACTOR
        if count > unrealistic_limit:
            checks.append(ValidationCheckResult(
                category="page_count", field="halaman_inti",
                status="skipped",
                message=(
                    f"Jumlah halaman terhitung tidak realistis "
                    f"({count} halaman > {unrealistic_limit}× batas maks). "
                    "Dokumen mungkin belum pernah dirender oleh Word "
                    "sehingga tidak memiliki penanda page break."
                ),
                skip_reason=f"hitung halaman tidak realistis: {count}",
            ))
            return issues, checks

        occurrences: list[dict] = []
        current_bab: str | None = start_text
        for i in range(start_idx, end_idx + 1):
            para = para_list[i]
            text = para.text.strip()
            if not text:
                continue
            if _heading_level_from_style(para.style) is None:
                continue
            pg = page_map.get(i, 1)
            if _BAB_RE.match(text.upper()):
                current_bab = text
            occurrences.append({
                "text"     : text[:100],
                "full_text": text,
                "style"    : para.style.name,
                "page"     : pg,
                "bab"      : current_bab,
                "para_idx" : i,
                "actual"   : f"halaman {pg}",
                "expected" : None,
            })

        if min_pages is not None:
            expected_str = f"{min_pages}–{maks} halaman"
        else:
            expected_str = f"≤ {maks} halaman"
        actual_str = f"{count} halaman"

        too_many = count > maks
        too_few  = min_pages is not None and count < min_pages

        if not too_many and not too_few:
            checks.append(ValidationCheckResult(
                category="page_count", field="halaman_inti",
                status="passed",
                message=(
                    f"Jumlah halaman inti {count} halaman "
                    f"(halaman {start_page}–{end_page}): "
                    f"sesuai batas {expected_str}"
                ),
                expected=expected_str,
                actual=actual_str,
                occurrences=occurrences,
            ))
        else:
            if too_many:
                msg = (
                    f"Jumlah halaman inti {count} halaman "
                    f"(halaman {start_page}–{end_page}) "
                    f"melebihi batas maksimum {maks} halaman."
                )
            else:
                msg = (
                    f"Jumlah halaman inti {count} halaman "
                    f"(halaman {start_page}–{end_page}) "
                    f"kurang dari batas minimum {min_pages} halaman."
                )
            issues.append(ValidationIssue(
                category="page_count", field="halaman_inti",
                severity="error", message=msg,
                expected=expected_str,
                actual=actual_str,
                occurrences=occurrences,
            ))
            checks.append(ValidationCheckResult(
                category="page_count", field="halaman_inti",
                status="failed", message=msg,
                expected=expected_str,
                actual=actual_str,
                occurrences=occurrences,
            ))

    except Exception as exc:
        checks.append(ValidationCheckResult(
            category="page_count", field="halaman_inti",
            status="skipped",
            message=f"Pengecekan jumlah halaman inti dilewati: {exc}",
            skip_reason=str(exc),
        ))

    return issues, checks
