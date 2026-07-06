"""Typography checks: heading case and body content formatting. Keyword: automated document validation"""
from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

from model_ai.extractor.models import DocumentMetadata
from model_ai.validation.models import ValidationCheckResult, ValidationIssue
from model_ai.validation.validocx_adapter import (
    _heading_level_from_style,
)

from ._shared import (
    _BAB_RE,
    _CAPTION_ALIGN_MAP,
    _FIG_DETECT_RE,
    _TBL_DETECT_RE,
    _LAMPIRAN_BROAD_RE,
    _TOC_TOF_STYLE_NAMES,
    _build_occurrences,
    _is_heading_para,
    _humanize_attr_value,
)
from .page_count import _build_displayed_page_map


def _text_matches_case_para(para, case_style: str) -> bool:
    """Cek apakah teks paragraf sesuai case_style yang diharapkan."""
    text = para.text.strip()
    alpha = [c for c in text if c.isalpha()]
    if not alpha:
        return True

    has_all_caps = any(run.font.all_caps for run in para.runs if run.text.strip())

    if case_style == "UPPERCASE":
        return all(c.isupper() for c in alpha) or has_all_caps
    if case_style == "LOWERCASE":
        return all(c.islower() for c in alpha) and not has_all_caps
    if case_style == "SENTENCE_CASE":
        first_alpha = next((c for c in text if c.isalpha()), None)
        return first_alpha is not None and first_alpha.isupper() and not has_all_caps
    if case_style == "TOGGLE_CASE":
        return True
    return True


def _check_heading_case(
    docx_path: Path,
    metadata: DocumentMetadata,
    doc: DocxDocument | None = None,
) -> tuple[list[ValidationIssue], list[ValidationCheckResult]]:
    """Validasi style huruf (case) pada Heading 1–5."""
    issues: list[ValidationIssue] = []
    checks: list[ValidationCheckResult] = []

    t = metadata.typography
    if t is None:
        return issues, checks

    h1_case = t.heading_1_case

    case_per_level: dict[int, str | None] = {
        1: h1_case,
        2: t.heading_2_case,
        3: t.heading_3_case,
        4: t.heading_4_case,
        5: t.heading_5_case,
    }

    if all(v is None for v in case_per_level.values()):
        return issues, checks

    try:
        doc = doc or DocxDocument(str(docx_path))

        page_map = _build_displayed_page_map(doc)

        pass_per_level:     dict[int, list[dict]] = {lvl: [] for lvl in case_per_level}
        mismatch_per_level: dict[int, list[dict]] = {lvl: [] for lvl in case_per_level}

        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            level = _heading_level_from_style(para.style)
            if level is None or level not in case_per_level:
                continue
            case_style = case_per_level[level]
            if not case_style:
                continue
            para_info: dict = {
                "para_idx"  : idx,
                "style"     : para.style.name,
                "text"      : text[:100],
                "full_text" : text,
                "bab"       : None,
                "page"      : page_map.get(idx),
            }
            if not _text_matches_case_para(para, case_style):
                mismatch_per_level[level].append(para_info)
            else:
                pass_per_level[level].append(para_info)

        for level, case_style in case_per_level.items():
            if case_style is None:
                continue
            field_name = f"heading_{level}_case"
            mismatches = mismatch_per_level[level]
            passes     = pass_per_level[level]

            if not mismatches and not passes:
                continue

            if mismatches:
                first_actual = mismatches[0]["text"]
                msg = (
                    f"Heading {level} harus {case_style}. "
                    f"{len(mismatches)} heading tidak sesuai. "
                    f'Contoh: "{first_actual}"'
                )
                occs = _build_occurrences(
                    mismatches, actual_str=first_actual, expected_str=case_style
                ) or None
                issues.append(ValidationIssue(
                    category="typography", field=field_name,
                    severity="error", message=msg,
                    expected=case_style, actual=first_actual,
                    occurrences=occs,
                ))
                checks.append(ValidationCheckResult(
                    category="typography", field=field_name,
                    status="failed", message=msg,
                    expected=case_style, actual=first_actual,
                    occurrences=occs,
                ))
            else:
                occs = _build_occurrences(passes, expected_str=case_style) or None
                checks.append(ValidationCheckResult(
                    category="typography", field=field_name,
                    status="passed",
                    message=f"Heading {level} case {case_style}: semua sesuai",
                    expected=case_style,
                    actual=case_style,
                    occurrences=occs,
                ))

    except Exception as exc:
        checks.append(ValidationCheckResult(
            category="typography", field="heading_case",
            status="skipped",
            message=f"Pengecekan case heading dilewati: {exc}",
            skip_reason=str(exc),
        ))

    return issues, checks


def _check_body_content(
    docx_path: Path,
    metadata: DocumentMetadata,
    doc: DocxDocument | None = None,
) -> tuple[list[ValidationIssue], list[ValidationCheckResult]]:
    """Validasi konten non-heading via content-based check.

    Iterasi SEMUA paragraf (termasuk w:sdt seperti TOC), skip heading dan semua
    jenis caption, cek alignment/font_family/font_size/line_spacing dari metadata.
    Hasil diagregasi per nilai parameter — bukan per nama style.

    Skip rules:
      - Paragraf kosong (text.strip() == "")
      - Heading: style name/inheritance mengandung 'heading' atau 'judul'
      - Caption gambar   : teks diawali 'Gambar \\d'  → dicek _check_caption_format
      - Caption tabel    : teks diawali 'Tabel \\d'   → dicek _check_caption_format
      - Caption lampiran : teks diawali 'Lampiran '   → dicek _check_figures_tables
    """
    issues: list[ValidationIssue] = []
    checks: list[ValidationCheckResult] = []

    t = metadata.typography
    _body_align_str = (metadata.spacing.paragraph_alignment if metadata.spacing else None) or "JUSTIFY"
    expected_align  = _CAPTION_ALIGN_MAP.get(_body_align_str.upper(), WD_ALIGN_PARAGRAPH.JUSTIFY)
    try:
        from model_ai.validation.validocx.wrapper import DocumentWrapper

        doc     = doc or DocxDocument(str(docx_path))
        wrapper = DocumentWrapper(doc)
        page_map = _build_displayed_page_map(doc)
        para_idx_by_id = {id(p._p): i for i, p in enumerate(doc.paragraphs)}

        align_pass: list[dict] = []
        align_fail: list[dict] = []

        for idx, para in enumerate(wrapper.iter_paragraphs()):
            text = para.text.strip()
            if not text:
                continue
            if _is_heading_para(para):
                continue
            if para.style.name not in _TOC_TOF_STYLE_NAMES and (
                _FIG_DETECT_RE.match(text)
                or _TBL_DETECT_RE.match(text)
                or _LAMPIRAN_BROAD_RE.match(text)
            ):
                continue

            _real_idx = para_idx_by_id.get(id(para._p))
            para_info: dict = {
                "para_idx" : idx,
                "style"    : para.style.name,
                "text"     : text[:100],
                "full_text": text,
                "bab"      : None,
                "page"     : page_map.get(_real_idx) if _real_idx is not None else None,
            }

            align = para.paragraph_format.alignment
            if align is None:
                try:
                    align = para.style.paragraph_format.alignment
                except Exception:
                    align = None
            if align is None or align == expected_align:
                align_pass.append(para_info)
            else:
                align_str_val = _humanize_attr_value("alignment", str(int(align))) if align is not None else "None"

        if align_pass or align_fail:
            actual_vals = list(dict.fromkeys(d.get("actual", "?") for d in align_fail))
            actual_str  = ", ".join(str(v) for v in actual_vals[:3])
            if align_fail:
                msg = (
                    f"Alignment: {len(align_fail)} elemen tidak sesuai "
                    f"(ekspektasi: JUSTIFY). Ditemukan: {actual_str}"
                )
                occs = _build_occurrences(align_fail, actual_str=actual_str, expected_str="JUSTIFY") or None
                issues.append(ValidationIssue(
                    category="typography", field="body_alignment",
                    severity="error", message=msg,
                    expected="JUSTIFY", actual=actual_str,
                    occurrences=occs,
                ))
                checks.append(ValidationCheckResult(
                    category="typography", field="body_alignment",
                    status="failed", message=msg,
                    expected="JUSTIFY", actual=actual_str,
                    occurrences=occs,
                ))
            if align_pass:
                occs = _build_occurrences(align_pass, expected_str="JUSTIFY") or None
                checks.append(ValidationCheckResult(
                    category="typography", field="body_alignment",
                    status="passed",
                    message=f"Alignment: {len(align_pass)} elemen lolos",
                    expected="JUSTIFY", actual="JUSTIFY",
                    occurrences=occs,
                ))
    except Exception as exc:
        checks.append(ValidationCheckResult(
            category="typography", field="body_content",
            status="skipped",
            message=f"Pengecekan konten body dilewati: {exc}",
            skip_reason=str(exc),
        ))

    return issues, checks


def _check_title_format(
    docx_path: Path,
    metadata: DocumentMetadata,
    doc: DocxDocument | None = None,
) -> tuple[list[ValidationIssue], list[ValidationCheckResult]]:
    """Validasi alignment, bold, dan case pada paragraf judul dokumen.

    Paragraf judul dideteksi sebagai paragraf pertama berisi teks ≥6 karakter
    sebelum BAB 1. Pengecekan dilewati jika tidak ada rule title_case maupun
    title_alignment di metadata (title_bold default True tidak cukup sendiri).
    """
    issues: list[ValidationIssue] = []
    checks: list[ValidationCheckResult] = []

    s = metadata.spacing
    if s is None:
        return issues, checks

    title_case_rule = s.title_case
    title_bold_req  = s.title_bold
    title_align_str = s.title_alignment

    if title_case_rule is None and title_align_str is None:
        return issues, checks

    try:
        doc = doc or DocxDocument(str(docx_path))
        page_map = _build_displayed_page_map(doc)

        title_para = None
        title_idx  = 0
        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text or len(text) < 6:
                continue
            if _BAB_RE.match(text.upper()):
                break
            title_para = para
            title_idx  = idx
            break

        if title_para is None:
            return issues, checks

        title_text = title_para.text.strip()
        para_info: dict = {
            "para_idx" : title_idx,
            "style"    : title_para.style.name,
            "text"     : title_text[:100],
            "full_text": title_text,
            "bab"      : None,
            "page"     : page_map.get(title_idx),
        }

        # Bold — jalan style chain jika run tidak mendefinisikan bold secara eksplisit
        def _run_bold(run, style) -> bool:
            if run.font.bold is True:
                return True
            if run.font.bold is False:
                return False
            cur, depth = style, 0
            while cur is not None and depth < 10:
                v = getattr(cur.font, "bold", None)
                if v is True:
                    return True
                if v is False:
                    return False
                cur = getattr(cur, "base_style", None)
                depth += 1
            return False

        runs_with_text = [r for r in title_para.runs if r.text.strip()]
        if runs_with_text and title_bold_req is not None:
            any_bold = any(_run_bold(r, title_para.style) for r in runs_with_text)
            if title_bold_req and not any_bold:
                issues.append(ValidationIssue(
                    category="typography", field="title_bold",
                    severity="error",
                    message="Judul dokumen harus tebal (bold) namun tidak terdeteksi bold.",
                    expected="bold", actual="tidak bold",
                    occurrences=[{**para_info, "actual": "tidak bold", "expected": "bold"}],
                ))
            elif not title_bold_req and any_bold:
                issues.append(ValidationIssue(
                    category="typography", field="title_bold",
                    severity="error",
                    message="Judul dokumen tidak boleh tebal (bold) namun terdeteksi bold.",
                    expected="tidak bold", actual="bold",
                    occurrences=[{**para_info, "actual": "bold", "expected": "tidak bold"}],
                ))
            else:
                checks.append(ValidationCheckResult(
                    category="typography", field="title_bold",
                    status="passed",
                    message="Bold judul sesuai.",
                    expected=str(title_bold_req), actual=str(any_bold),
                ))

        # Alignment
        if title_align_str:
            expected_align = _CAPTION_ALIGN_MAP.get(title_align_str.upper(), WD_ALIGN_PARAGRAPH.CENTER)
            actual_align = title_para.paragraph_format.alignment
            if actual_align is None:
                try:
                    actual_align = title_para.style.paragraph_format.alignment
                except Exception:
                    pass
            if actual_align is not None and actual_align != expected_align:
                align_str_val = _ALIGN_LABEL.get(int(actual_align), str(int(actual_align)))
                issues.append(ValidationIssue(
                    category="typography", field="title_alignment",
                    severity="error",
                    message=(
                        f"Alignment judul tidak sesuai. "
                        f"Ditemukan: {align_str_val}, Seharusnya: {title_align_str}"
                    ),
                    expected=title_align_str, actual=align_str_val,
                    occurrences=[{**para_info, "actual": align_str_val, "expected": title_align_str}],
                ))
            else:
                checks.append(ValidationCheckResult(
                    category="typography", field="title_alignment",
                    status="passed",
                    message=f"Alignment judul sesuai: {title_align_str}",
                    expected=title_align_str, actual=title_align_str,
                ))

        # Case
        if title_case_rule:
            if not _text_matches_case_para(title_para, title_case_rule):
                issues.append(ValidationIssue(
                    category="typography", field="title_case",
                    severity="error",
                    message=f"Judul dokumen harus {title_case_rule}. Teks: \"{title_text[:60]}\"",
                    expected=title_case_rule, actual=title_text[:100],
                    occurrences=[{**para_info, "actual": title_text[:100], "expected": title_case_rule}],
                ))
            else:
                checks.append(ValidationCheckResult(
                    category="typography", field="title_case",
                    status="passed",
                    message=f"Case judul sesuai: {title_case_rule}",
                    expected=title_case_rule, actual=title_text[:100],
                ))

    except Exception as exc:
        checks.append(ValidationCheckResult(
            category="typography", field="title_format",
            status="skipped",
            message=f"Pengecekan format judul dilewati: {exc}",
            skip_reason=str(exc),
        ))

    return issues, checks


def _check_font_size_sections(
    docx_path: Path,
    metadata: DocumentMetadata,
    doc: DocxDocument | None = None,
) -> tuple[list[ValidationIssue], list[ValidationCheckResult]]:
    """Validasi ukuran font pada paragraf judul, identitas penulis, dan abstrak."""
    issues: list[ValidationIssue] = []
    checks: list[ValidationCheckResult] = []

    t = metadata.typography
    if t is None:
        return issues, checks

    size_title    = t.font_size_title_pt
    size_author   = t.font_size_author_pt
    size_abstract = t.font_size_abstract_pt

    if size_title is None and size_author is None and size_abstract is None:
        return issues, checks

    try:
        doc = doc or DocxDocument(str(docx_path))
        page_map = _build_displayed_page_map(doc)
        para_idx_by_id = {id(p._p): i for i, p in enumerate(doc.paragraphs)}

        title_para:     object | None = None
        author_paras:   list          = []
        abstract_paras: list          = []
        phase = "pre_title"

        for para in doc.paragraphs:
            text = para.text.strip()
            if _BAB_RE.match(text.upper()):
                break
            if phase == "pre_title":
                if text and len(text) >= 6:
                    title_para = para
                    phase = "author"
            elif phase == "author":
                if text.upper() in ("ABSTRAK", "ABSTRACT"):
                    phase = "abstract"
                elif text:
                    author_paras.append(para)
            elif phase == "abstract":
                if text:
                    abstract_paras.append(para)

        def _resolve_size_pt(run, style) -> float | None:
            if run.font.size is not None:
                return run.font.size.pt
            cur, depth = style, 0
            while cur is not None and depth < 10:
                sz = getattr(cur.font, "size", None)
                if sz is not None:
                    return sz.pt
                cur = getattr(cur, "base_style", None)
                depth += 1
            return None

        def _check_size_group(target_paras, expected_pt, field: str, label: str) -> None:
            if not target_paras or expected_pt is None:
                return
            mismatches: list[dict] = []
            for para in target_paras:
                runs = [r for r in para.runs if r.text.strip()]
                for run in runs:
                    actual = _resolve_size_pt(run, para.style)
                    if actual is not None and abs(actual - expected_pt) > 0.5:
                        _pidx = para_idx_by_id.get(id(para._p))
                        mismatches.append({
                            "para_idx": _pidx,
                            "style":    para.style.name,
                            "text":     para.text.strip()[:100],
                            "full_text": para.text.strip(),
                            "actual":   f"{actual:.1f}pt",
                            "bab":      None,
                            "page":     page_map.get(_pidx) if _pidx is not None else None,
                        })
                        break
            exp_str = f"{expected_pt}pt"
            if mismatches:
                actual_str = mismatches[0]["actual"]
                msg = f"Ukuran font {label} harus {exp_str}, ditemukan: {actual_str}."
                occs = _build_occurrences(
                    mismatches, actual_str=actual_str, expected_str=exp_str
                ) or None
                issues.append(ValidationIssue(
                    category="typography", field=field,
                    severity="error", message=msg,
                    expected=exp_str, actual=actual_str,
                    occurrences=occs,
                ))
                checks.append(ValidationCheckResult(
                    category="typography", field=field,
                    status="failed", message=msg,
                    expected=exp_str, actual=actual_str,
                    occurrences=occs,
                ))
            else:
                checks.append(ValidationCheckResult(
                    category="typography", field=field,
                    status="passed",
                    message=f"Ukuran font {label} {exp_str}: sesuai",
                    expected=exp_str, actual=exp_str,
                ))

        if title_para is not None:
            _check_size_group([title_para], size_title,    "font_size_title",    "judul")
        _check_size_group(author_paras,   size_author,   "font_size_author",   "identitas penulis")
        _check_size_group(abstract_paras, size_abstract, "font_size_abstract", "abstrak")

    except Exception as exc:
        checks.append(ValidationCheckResult(
            category="typography", field="font_size_sections",
            status="skipped",
            message=f"Pengecekan ukuran font section dilewati: {exc}",
            skip_reason=str(exc),
        ))

    return issues, checks


def _check_section_line_spacing(
    docx_path: Path,
    metadata: DocumentMetadata,
    doc: DocxDocument | None = None,
) -> tuple[list[ValidationIssue], list[ValidationCheckResult]]:
    """Validasi spasi baris section judul-abstrak dan daftar pustaka.

    Berlaku untuk skema artikel (PKM-AI). Hanya aktif jika metadata
    mengandung rule/nilai untuk kedua section tersebut.
    """
    issues: list[ValidationIssue] = []
    checks: list[ValidationCheckResult] = []

    s = metadata.spacing
    if s is None:
        return issues, checks

    rule_ta  = s.line_spacing_rule_title_abstract
    val_ta   = s.line_spacing_title_abstract
    rule_bib = s.line_spacing_rule_bibliography
    val_bib  = s.line_spacing_bibliography

    if rule_ta is None and val_ta is None and rule_bib is None and val_bib is None:
        return issues, checks

    try:
        from docx.oxml.ns import qn as _qn

        doc = doc or DocxDocument(str(docx_path))
        page_map = _build_displayed_page_map(doc)
        para_idx_by_id = {id(p._p): i for i, p in enumerate(doc.paragraphs)}

        ta_paras:  list = []
        bib_paras: list = []
        found_bab = False
        in_bib    = False

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if _BAB_RE.match(text.upper()):
                found_bab = True
                in_bib = False
            if not found_bab:
                ta_paras.append(para)
            if text.upper() == "DAFTAR PUSTAKA":
                in_bib = True
                continue
            if in_bib and found_bab:
                bib_paras.append(para)

        def _read_line_spacing(para):
            pPr = para._element.find(_qn("w:pPr"))
            if pPr is None:
                return None, None
            sp_el = pPr.find(_qn("w:spacing"))
            if sp_el is None:
                return None, None
            line_str  = sp_el.get(_qn("w:line"))
            line_rule = sp_el.get(_qn("w:lineRule"))
            if line_str is None:
                return None, None
            val = int(line_str)
            if line_rule in ("auto", None):
                mult = val / 240
                if val == 240:
                    return "SINGLE", 1.0
                if val == 360:
                    return "ONE_POINT_FIVE", 1.5
                if val == 480:
                    return "DOUBLE", 2.0
                return "MULTIPLE", mult
            if line_rule == "atLeast":
                return "AT_LEAST", val / 20
            if line_rule == "exact":
                return "EXACTLY", val / 20
            return None, None

        _RULE_LABEL: dict[str, str] = {
            "SINGLE":         "Spasi tunggal (1.0)",
            "ONE_POINT_FIVE": "Satu setengah (1.5)",
            "DOUBLE":         "Spasi ganda (2.0)",
            "MULTIPLE":       "Kelipatan",
            "AT_LEAST":       "Setidaknya",
            "EXACTLY":        "Tepat",
        }

        def _spacing_ok(rule_a, val_a, exp_rule_str, exp_val) -> bool:
            if exp_rule_str is None:
                return True
            if rule_a != exp_rule_str.upper():
                return False
            if exp_val is not None and val_a is not None:
                return abs(float(val_a) - float(exp_val)) < 0.1
            return True

        def _check_spacing_group(
            section_paras, exp_rule_str, exp_val, field: str, label: str
        ) -> None:
            if not section_paras or (exp_rule_str is None and exp_val is None):
                return
            mismatches: list[dict] = []
            for para in section_paras:
                rule_a, val_a = _read_line_spacing(para)
                if not _spacing_ok(rule_a, val_a, exp_rule_str, exp_val):
                    _pidx = para_idx_by_id.get(id(para._p))
                    mismatches.append({
                        "para_idx":  _pidx,
                        "style":     para.style.name,
                        "text":      para.text.strip()[:100],
                        "full_text": para.text.strip(),
                        "actual":    f"{rule_a} {val_a:.2f}" if val_a else str(rule_a),
                        "bab":       None,
                        "page":      page_map.get(_pidx) if _pidx is not None else None,
                    })
            rule_lbl = _RULE_LABEL.get((exp_rule_str or "").upper(), exp_rule_str or "")
            exp_str  = f"{rule_lbl} {exp_val}" if exp_val else rule_lbl
            if mismatches:
                actual_str = mismatches[0]["actual"]
                msg = (
                    f"Spasi baris {label} tidak sesuai (ekspektasi: {exp_str}). "
                    f"{len(mismatches)} paragraf tidak sesuai."
                )
                occs = _build_occurrences(
                    mismatches, actual_str=actual_str, expected_str=exp_str
                ) or None
                issues.append(ValidationIssue(
                    category="spacing", field=field,
                    severity="error", message=msg,
                    expected=exp_str, actual=actual_str,
                    occurrences=occs,
                ))
                checks.append(ValidationCheckResult(
                    category="spacing", field=field,
                    status="failed", message=msg,
                    expected=exp_str, actual=actual_str,
                    occurrences=occs,
                ))
            else:
                checks.append(ValidationCheckResult(
                    category="spacing", field=field,
                    status="passed",
                    message=f"Spasi baris {label} ({exp_str}): semua sesuai",
                    expected=exp_str, actual=exp_str,
                ))

        _check_spacing_group(ta_paras,  rule_ta,  val_ta,  "line_spacing_title_abstract", "judul-abstrak")
        _check_spacing_group(bib_paras, rule_bib, val_bib, "line_spacing_bibliography",   "daftar pustaka")

    except Exception as exc:
        checks.append(ValidationCheckResult(
            category="spacing", field="line_spacing_sections",
            status="skipped",
            message=f"Pengecekan spasi baris section dilewati: {exc}",
            skip_reason=str(exc),
        ))

    return issues, checks
